from pathlib import Path
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Bao_cao_DATN_TQAuto_LeDinhQuoc_2022603416.docx"
SCREENSHOTS = ROOT / ".screenshots"


def set_cell_text(cell, text, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_static_toc(doc):
    rows = [
        ("DANH MỤC CÁC TỪ VIẾT TẮT", "4"),
        ("DANH MỤC BẢNG BIỂU", "5"),
        ("DANH MỤC HÌNH ẢNH", "6"),
        ("MỞ ĐẦU", "7"),
        ("1. Lý do chọn đề tài", "7"),
        ("2. Mục tiêu của đề tài", "7"),
        ("3. Đối tượng và phạm vi nghiên cứu", "8"),
        ("4. Ý nghĩa khoa học và thực tiễn", "8"),
        ("5. Bố cục của báo cáo", "8"),
        ("CHƯƠNG 1. CƠ SỞ LÝ THUYẾT", "9"),
        ("1.1. Tổng quan về thương mại điện tử trong lĩnh vực ô tô", "9"),
        ("1.2. Công nghệ phát triển ứng dụng", "9"),
        ("1.3. Kiến trúc ứng dụng Next.js", "10"),
        ("1.4. Cơ sở dữ liệu MySQL", "11"),
        ("1.5. Bảo mật và phân quyền", "12"),
        ("CHƯƠNG 2. KHẢO SÁT, PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "13"),
        ("2.1. Khảo sát hiện trạng", "13"),
        ("2.2. Tác nhân của hệ thống", "13"),
        ("2.3. Yêu cầu chức năng", "14"),
        ("2.4. Yêu cầu phi chức năng", "15"),
        ("2.5. Sơ đồ Use Case", "16"),
        ("2.6. Đặc tả & Phân tích chi tiết ca sử dụng", "18"),
        ("2.7. Thiết kế cơ sở dữ liệu chi tiết", "29"),
        ("2.8. Thiết kế giao diện tổng quan", "34"),
        ("CHƯƠNG 3. KẾT QUẢ CÀI ĐẶT VÀ KIỂM THỬ HỆ THỐNG", "35"),
        ("3.1. Sơ đồ điều hướng màn hình", "35"),
        ("3.2. Thiết kế chi tiết màn hình", "36"),
        ("3.3. Kiểm thử hệ thống", "38"),
        ("3.4. Đánh giá kết quả", "42"),
        ("KẾT LUẬN", "43"),
        ("HƯỚNG PHÁT TRIỂN", "43"),
        ("TÀI LIỆU THAM KHẢO", "43"),
    ]
    for title, page in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run(title)
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)
        if title.isupper() or title.startswith("CHƯƠNG"):
            r.bold = True
        dots = "." * max(6, 70 - len(title))
        r2 = p.add_run(f" {dots} ")
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(13)
        r3 = p.add_run(page)
        r3.font.name = "Times New Roman"
        r3.font.size = Pt(13)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)


def transform_desc(original_desc):
    desc_dict = {k: v for k, v in original_desc}
    name = desc_dict.get("Tên ca sử dụng") or desc_dict.get("Tên use case") or ""
    mota = desc_dict.get("Mô tả vắn tắt") or ""
    luong_chinh = desc_dict.get("Luồng sự kiện chính") or ""
    luong_re_nhanh = desc_dict.get("Luồng rẽ nhánh") or ""
    
    luong_su_kien_text = luong_chinh
    if luong_re_nhanh:
        db_fail_flow = "3. Tại tất cả các bước trong luồng cơ bản nếu hệ thống không thể kết nối được với cơ sở dữ liệu thì sẽ hiển thị một thông báo lỗi. Use case kết thúc."
        luong_su_kien_text += f"\n\nLuồng rẽ nhánh:\n{luong_re_nhanh}\n{db_fail_flow}"
    else:
        db_fail_flow = "1. Tại tất cả các bước trong luồng cơ bản nếu hệ thống không thể kết nối được với cơ sở dữ liệu thì sẽ hiển thị một thông báo lỗi. Use case kết thúc."
        luong_su_kien_text += f"\n\nLuồng rẽ nhánh:\n{db_fail_flow}"
        
    tac_nhan = desc_dict.get("Tác nhân", "").lower()
    if "quản trị viên" in tac_nhan or "admin" in tac_nhan:
        special_req = "Người dùng phải có vai trò là người quản trị"
    elif "nhân viên" in tac_nhan or "staff" in tac_nhan:
        special_req = "Người dùng phải có vai trò là nhân viên kinh doanh hoặc quản trị viên"
    else:
        special_req = "Không có"
        
    tiendieukien = desc_dict.get("Tiền điều kiện") or "Không có"
    haudieukien = desc_dict.get("Hậu điều kiện") or "Không có"
    
    new_desc = [
        ["Tên use case", name],
        ["Mô tả vắn tắt", mota],
        ["", "Luồng cơ bản:"],
        ["Luồng sự kiện", luong_su_kien_text],
        ["Các yêu cầu đặc biệt", special_req],
        ["Tiền điều kiện", tiendieukien],
        ["Hậu điều kiện", haudieukien],
        ["Điểm mở rộng", "Không có"]
    ]
    return new_desc


def add_usecase_table(doc, rows, widths=None):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    set_table_borders(table)
    
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        key, val = row
        
        align_key = WD_ALIGN_PARAGRAPH.LEFT
        align_val = WD_ALIGN_PARAGRAPH.LEFT
        
        is_luong_co_ban = (key == "" and val == "Luồng cơ bản:")
        
        if is_luong_co_ban:
            set_cell_text(cells[0], "", align=align_key)
            set_cell_text(cells[1], val, bold=True, align=align_val)
        else:
            set_cell_text(cells[0], key, bold=True, align=align_key)
            set_cell_text(cells[1], val, align=align_val)
            
        if widths:
            set_cell_width(cells[0], widths[0])
            set_cell_width(cells[1], widths[1])
            
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr[idx], "E0E0E0") # Grayscale header fill
        if widths:
            set_cell_width(hdr[idx], widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            val_str = str(val)
            align = WD_ALIGN_PARAGRAPH.CENTER if len(val_str) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], val_str, align=align)
            if widths:
                set_cell_width(cells[idx], widths[idx])
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def add_image(doc, filename, caption):
    path = SCREENSHOTS / filename
    if not path.exists():
        p_err = doc.add_paragraph()
        p_err.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_err.add_run(f"[Hình ảnh bị thiếu: {filename}]").font.color.rgb = RGBColor(255, 0, 0)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.1))
    add_caption(doc, caption)


def add_cover(doc):
    for text, size, bold in [
        ("BỘ CÔNG THƯƠNG", 13, True),
        ("ĐẠI HỌC CÔNG NGHIỆP HÀ NỘI", 13, True),
        ("---------------------------", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ĐỒ ÁN TỐT NGHIỆP")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NGÀNH KHOA HỌC MÁY TÍNH")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TÊN ĐỀ TÀI")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("XÂY DỰNG WEBSITE BÁN Ô TÔ TRỰC TUYẾN\nDÀNH CHO SHOWROOM TQ AUTO")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)

    doc.add_paragraph()
    doc.add_paragraph()
    rows = [
        ("GVHD:", "TS. Phạm Văn Hiệp"),
        ("Sinh viên:", "Lê Đình Quốc"),
        ("Mã sinh viên:", "2022603416"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(cells[1], right, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_width(cells[0], 2500)
        set_cell_width(cells[1], 4200)

    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hà Nội - Năm 2026")
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    doc.add_page_break()


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(13)
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.space_after = Pt(6)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
def main():
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    # List of all 19 use cases
    uc_specs = [
        {
            "name": "Đăng ký tài khoản",
            "seq_img": "seq_register.png", "class_img": "class_register.png",
            "desc": [
                ["Tên ca sử dụng", "Đăng ký tài khoản"],
                ["Tác nhân", "Khách vãng lai"],
                ["Mô tả vắn tắt", "Cho phép khách vãng lai tạo tài khoản thành viên mới để sử dụng các tính năng cá nhân hóa."],
                ["Tiền điều kiện", "Khách hàng chưa đăng nhập và đang ở trang đăng ký."],
                ["Hậu điều kiện", "Tài khoản khách hàng mới được lưu vào bảng users ở trạng thái hoạt động (active)."],
                ["Luồng sự kiện chính", "1. Khách hàng click nút 'Đăng ký'.\n2. Hệ thống hiển thị biểu mẫu đăng ký.\n3. Khách hàng nhập Họ tên, Email, Số điện thoại và Mật khẩu.\n4. Hệ thống kiểm tra tính hợp lệ và đối chiếu email chưa tồn tại trong bảng users.\n5. Mã hóa mật khẩu bằng SHA-256 và INSERT bản ghi mới.\n6. Trả về thông báo thành công và chuyển hướng về trang chủ."],
                ["Luồng rẽ nhánh", "4a. Nếu email đã tồn tại: Hệ thống hiển thị thông báo lỗi 'Email này đã được sử dụng'.\n3a. Nếu thông tin thiếu hoặc SĐT sai định dạng: Báo lỗi đỏ dưới trường tương ứng."]
            ]
        },
        {
            "name": "Đăng nhập",
            "seq_img": "seq_login.png", "class_img": "class_login.png",
            "desc": [
                ["Tên ca sử dụng", "Đăng nhập"],
                ["Tác nhân", "Khách hàng, Nhân viên, Quản trị viên"],
                ["Mô tả vắn tắt", "Xác thực danh tính người dùng và cấp quyền truy cập các chức năng tương ứng với vai trò."],
                ["Tiền điều kiện", "Người dùng đã có tài khoản đang ở trạng thái hoạt động."],
                ["Hậu điều kiện", "Khởi tạo mã token JWT lưu trong cookie HttpOnly để duy trì phiên."],
                ["Luồng sự kiện chính", "1. Người dùng nhập Email và Mật khẩu, bấm nút 'Đăng nhập'.\n2. Hệ thống truy vấn thông tin tài khoản trong bảng users theo Email.\n3. Đối chiếu mật khẩu băm và kiểm tra trạng thái hoạt động (status = active).\n4. Sinh mã token JWT và trả về trình duyệt.\n5. Chuyển hướng người dùng về trang chủ hoặc trang quản trị theo vai trò."],
                ["Luồng rẽ nhánh", "2a. Nếu Email không tồn tại hoặc sai mật khẩu: Báo lỗi 'Email hoặc mật khẩu không chính xác'.\n3a. Nếu tài khoản bị khóa (blocked): Báo lỗi 'Tài khoản của bạn đã bị khóa'."]
            ]
        },
        {
            "name": "Xem chi tiết xe",
            "seq_img": "seq_view_car.png", "class_img": "class_view_car.png",
            "desc": [
                ["Tên ca sử dụng", "Xem chi tiết xe"],
                ["Tác nhân", "Khách vãng lai, Khách hàng"],
                ["Mô tả vắn tắt", "Cho phép người dùng xem thông số kỹ thuật chi tiết, album ảnh và đánh giá của một chiếc xe cụ thể."],
                ["Tiền điều kiện", "Chiếc xe được chọn phải tồn tại và đang mở bán (status != hidden)."],
                ["Hậu điều kiện", "Hiển thị đầy đủ thông tin chi tiết xe, lượt xem (views) của xe tăng 1."],
                ["Luồng sự kiện chính", "1. Khách hàng click vào thẻ xe trong danh sách.\n2. Hệ thống gửi request lấy chi tiết xe.\n3. Server SELECT dữ liệu từ bảng cars và car_images.\n4. Thực hiện lệnh UPDATE cars SET views = views + 1.\n5. Trả về thông tin và render giao diện chi tiết xe."],
                ["Luồng rẽ nhánh", "3a. Nếu ID xe không tồn tại hoặc đã bị ẩn: Hệ thống chuyển hướng người dùng sang trang lỗi 404."]
            ]
        },
        {
            "name": "Xem tin tức",
            "seq_img": "seq_view_posts.png", "class_img": "class_view_posts.png",
            "desc": [
                ["Tên ca sử dụng", "Xem tin tức"],
                ["Tác nhân", "Khách vãng lai, Khách hàng"],
                ["Mô tả vắn tắt", "Người dùng đọc bài viết đánh giá xe, tin tức thị trường và thông tin từ showroom."],
                ["Tiền điều kiện", "Bài viết phải ở trạng thái xuất bản (published = 1)."],
                ["Hậu điều kiện", "Hiển thị bài viết tin tức đầy đủ, lượt xem tin tức tăng 1."],
                ["Luồng sự kiện chính", "1. Khách hàng click chọn danh mục Tin tức.\n2. Hệ thống SELECT danh sách bài viết từ bảng posts và render danh sách.\n3. Khách hàng click chọn một bài viết cụ thể.\n4. Hệ thống SELECT nội dung chi tiết của bài viết theo slug.\n5. UPDATE posts SET views = views + 1.\n6. Hiển thị nội dung bài viết tin tức đầy đủ."],
                ["Luồng rẽ nhánh", "4a. Nếu bài viết không tồn tại hoặc chưa xuất bản: Hiển thị trang báo lỗi 404."]
            ]
        },
        {
            "name": "Đặt lịch xem xe & lái thử",
            "seq_img": "seq_appointment.png", "class_img": "class_appointment.png",
            "desc": [
                ["Tên ca sử dụng", "Đặt lịch xem xe & lái thử"],
                ["Tác nhân", "Khách hàng"],
                ["Mô tả vắn tắt", "Khách hàng đăng ký ngày giờ đến showroom lái thử hoặc xem trực tiếp chiếc xe."],
                ["Tiền điều kiện", "Khách hàng đang ở trang chi tiết xe, khung giờ hẹn phải trống (< 3 lịch hẹn cùng lúc)."],
                ["Hậu điều kiện", "Lịch hẹn được lưu ở trạng thái pending, tạo thông báo hệ thống và gửi email xác nhận."],
                ["Luồng sự kiện chính", "1. Khách hàng mở biểu mẫu đặt lịch, chọn Ngày hẹn, Giờ hẹn và điền thông tin.\n2. Hệ thống kiểm tra tải trọng khung giờ (tối đa 3 lịch hẹn/giờ) và giờ làm việc.\n3. INSERT bản ghi vào bảng appointments với trạng thái pending.\n4. Tạo thông báo mới cho nhân viên và admin phụ trách.\n5. Gửi email xác nhận tự động cho khách hàng thông qua dịch vụ SMTP."],
                ["Luồng rẽ nhánh", "2a. Nếu khung giờ đã đầy hoặc ngoài giờ hành chính: Báo lỗi và yêu cầu chọn giờ khác.\n1a. Nếu ngày hẹn ở quá khứ hoặc thông tin trống: Hiển thị nhắc nhở báo đỏ dưới form."]
            ]
        },
        {
            "name": "Cập nhật hồ sơ cá nhân",
            "seq_img": "seq_update_profile.png", "class_img": "class_update_profile.png",
            "desc": [
                ["Tên ca sử dụng", "Cập nhật hồ sơ cá nhân"],
                ["Tác nhân", "Khách hàng"],
                ["Mô tả vắn tắt", "Cho phép thành viên cập nhật thông tin cá nhân như họ tên, số điện thoại, mật khẩu hoặc ảnh đại diện."],
                ["Tiền điều kiện", "Khách hàng đã đăng nhập thành công vào tài khoản."],
                ["Hậu điều kiện", "Thông tin tài khoản trong bảng users được cập nhật, session được làm mới."],
                ["Luồng sự kiện chính", "1. Khách hàng truy cập trang quản lý hồ sơ.\n2. Nhập thông tin thay đổi hoặc tải lên ảnh đại diện mới.\n3. Bấm nút 'Lưu cập nhật'.\n4. Hệ thống kiểm tra tính hợp lệ của số điện thoại và tải ảnh lên server.\n5. Thực hiện UPDATE users SET full_name, phone, avatar WHERE id.\n6. Trả về kết quả thành công và revalidate giao diện."],
                ["Luồng rẽ nhánh", "4a. Nếu số điện thoại không hợp lệ hoặc upload ảnh lỗi: Báo lỗi đỏ và giữ nguyên dữ liệu cũ."]
            ]
        },
        {
            "name": "Liên hệ",
            "seq_img": "seq_submit_consultation.png", "class_img": "class_submit_consultation.png",
            "desc": [
                ["Tên ca sử dụng", "Liên hệ"],
                ["Tác nhân", "Khách vãng lai, Khách hàng"],
                ["Mô tả vắn tắt", "Người dùng gửi yêu cầu tư vấn chung hoặc lời nhắn hỗ trợ về dịch vụ của showroom."],
                ["Tiền điều kiện", "Truy cập trang liên hệ hoặc chân trang."],
                ["Hậu điều kiện", "Yêu cầu lưu vào contact_requests, tạo lead mới trong CRM, gửi email SMTP báo nhận thông tin."],
                ["Luồng sự kiện chính", "1. Người dùng điền Họ tên, Email, Số điện thoại và nội dung tin nhắn liên hệ.\n2. Bấm nút 'Gửi liên hệ'.\n3. Hệ thống validate thông tin.\n4. INSERT yêu cầu vào bảng contact_requests.\n5. Tạo hoặc cập nhật hồ sơ khách hàng tiềm năng trong bảng customers CRM.\n6. Gửi email xác nhận qua Nodemailer và báo gửi thành công."],
                ["Luồng rẽ nhánh", "3a. Nếu thiếu các trường bắt buộc hoặc email sai định dạng: Chặn submit và báo lỗi dưới form."]
            ]
        },
        {
            "name": "Theo dõi lịch hẹn cá nhân",
            "seq_img": "seq_view_bookings.png", "class_img": "class_view_bookings.png",
            "desc": [
                ["Tên ca sử dụng", "Theo dõi lịch hẹn cá nhân"],
                ["Tác nhân", "Khách hàng"],
                ["Mô tả vắn tắt", "Cho phép khách hàng xem danh sách và trạng thái xử lý các lịch hẹn xem xe của chính mình."],
                ["Tiền điều kiện", "Khách hàng đã đăng nhập thành công."],
                ["Hậu điều kiện", "Hiển thị danh sách lịch hẹn kèm thông tin xe và trạng thái."],
                ["Luồng sự kiện chính", "1. Khách hàng vào mục 'Lịch hẹn của tôi' trong trang cá nhân.\n2. Hệ thống gọi API lấy danh sách lịch hẹn.\n3. Server SELECT từ bảng appointments JOIN cars WHERE user_id.\n4. Trả về mảng dữ liệu lịch hẹn của khách hàng.\n5. Client hiển thị danh sách lịch hẹn dạng timeline kèm trạng thái (chờ duyệt, đã duyệt, hoàn thành, hủy)."]
            ]
        },
        {
            "name": "Chat trực tuyến",
            "seq_img": "seq_customer_chat.png", "class_img": "class_customer_chat.png",
            "desc": [
                ["Tên ca sử dụng", "Chat trực tuyến"],
                ["Tác nhân", "Khách hàng, Nhân viên"],
                ["Mô tả vắn tắt", "Hỗ trợ khách hàng nhắn tin trao đổi trực tiếp thời gian thực với nhân viên showroom."],
                ["Tiền điều kiện", "Người dùng mở widget chat trên website."],
                ["Hậu điều kiện", "Tin nhắn được lưu vào chat_messages, hiển thị tức thời cho cả hai bên."],
                ["Luồng sự kiện chính", "1. Khách hàng mở widget chat, hệ thống nhận diện theo session_id.\n2. Khách hàng nhập tin nhắn và bấm Gửi.\n3. INSERT tin nhắn mới vào chat_messages với role = customer.\n4. Hệ thống báo hiệu có tin nhắn mới cho dashboard nhân viên.\n5. Nhân viên phản hồi, INSERT tin nhắn với role = staff.\n6. Cập nhật nội dung chat thời gian thực trên màn hình."],
                ["Luồng rẽ nhánh", "3a. Nếu lỗi kết nối mạng: Hiển thị biểu tượng chấm than đỏ bên cạnh tin nhắn báo gửi thất bại."]
            ]
        },
        {
            "name": "Gửi yêu cầu tư vấn",
            "seq_img": "seq_submit_lead.png", "class_img": "class_submit_lead.png",
            "desc": [
                ["Tên ca sử dụng", "Gửi yêu cầu tư vấn"],
                ["Tác nhân", "Khách hàng"],
                ["Mô tả vắn tắt", "Khách hàng gửi biểu mẫu yêu cầu tư vấn chuyên sâu về xe đang xem hoặc thủ tục trả góp."],
                ["Tiền điều kiện", "Khách hàng đang ở trang chi tiết của xe."],
                ["Hậu điều kiện", "Bản ghi yêu cầu tư vấn được lưu trong contact_requests và hồ sơ CRM được cập nhật."],
                ["Luồng sự kiện chính", "1. Khách hàng nhập Họ tên, Số điện thoại, Email, chọn loại tư vấn (Lái thử, Trả góp...) và gửi.\n2. Hệ thống kiểm tra dữ liệu.\n3. Server INSERT thông tin vào bảng contact_requests.\n4. Tạo hoặc cập nhật hồ sơ khách hàng tiềm năng trong bảng customers CRM ở giai đoạn new_lead.\n5. Bắn thông báo nội bộ cho nhân viên quản trị.\n6. Hiển thị popup thông báo gửi thành công."],
                ["Luồng rẽ nhánh", "2a. Nếu thiếu thông tin liên lạc: Hệ thống nhắc nhở điền đầy đủ số điện thoại và email."]
            ]
        },
        {
            "name": "Quản lý khách hàng CRM",
            "seq_img": "seq_crm.png", "class_img": "class_crm.png",
            "desc": [
                ["Tên ca sử dụng", "Quản lý khách hàng CRM"],
                ["Tác nhân", "Nhân viên, Quản trị viên"],
                ["Mô tả vắn tắt", "Quản lý hồ sơ khách hàng, cập nhật giai đoạn phễu chăm sóc và ghi chú lịch sử tư vấn."],
                ["Tiền điều kiện", "Nhân viên đã đăng nhập và có quyền truy cập CRM."],
                ["Hậu điều kiện", "Giai đoạn chuyển đổi được cập nhật trong customers, ghi chú lưu vào customer_notes."],
                ["Luồng sự kiện chính", "1. Nhân viên truy cập danh sách khách hàng trong CRM.\n2. Chọn một khách hàng cụ thể để xem chi tiết timeline chăm sóc.\n3. Nhân viên nhập ghi chú cuộc gọi tư vấn và chọn giai đoạn mới (ví dụ: đang đàm phán).\n4. Bấm 'Lưu cập nhật'.\n5. Server thực hiện UPDATE customers SET stage = 'negotiating' và INSERT vào customer_notes.\n6. Timeline chăm sóc được làm mới với thông tin cập nhật."],
                ["Luồng rẽ nhánh", "3a. Ghi chú trống: Hệ thống hiển thị cảnh báo yêu cầu nhập nội dung trao đổi."]
            ]
        },
        {
            "name": "Quản lý lịch hẹn",
            "seq_img": "seq_manage_appts.png", "class_img": "class_manage_appts.png",
            "desc": [
                ["Tên ca sử dụng", "Quản lý lịch hẹn"],
                ["Tác nhân", "Nhân viên, Quản trị viên"],
                ["Mô tả vắn tắt", "Xem, xác nhận, hoàn thành hoặc hủy bỏ lịch hẹn xem xe của khách hàng."],
                ["Tiền điều kiện", "Nhân viên đã đăng nhập và được phân công duyệt lịch hẹn."],
                ["Hậu điều kiện", "Trạng thái lịch hẹn được cập nhật trong database, email thông báo được gửi cho khách."],
                ["Luồng sự kiện chính", "1. Nhân viên mở danh sách lịch hẹn đang chờ duyệt.\n2. Kiểm tra thông tin xe và thời gian hẹn, bấm nút 'Xác nhận'.\n3. Server UPDATE appointments SET status = 'confirmed'.\n4. Hệ thống tự động gửi email thông báo xác nhận thời gian hẹn cho khách hàng.\n5. Cập nhật bảng dữ liệu lịch hẹn hiển thị trạng thái mới."],
                ["Luồng rẽ nhánh", "2a. Nếu nhân viên bấm hủy lịch: Trạng thái chuyển sang 'cancelled' và yêu cầu nhập lý do hủy gửi cho khách."]
            ]
        },
        {
            "name": "Quản lý bài viết",
            "seq_img": "seq_manage_posts.png", "class_img": "class_manage_posts.png",
            "desc": [
                ["Tên ca sử dụng", "Quản lý bài viết"],
                ["Tác nhân", "Nhân viên, Quản trị viên"],
                ["Mô tả vắn tắt", "Cho phép đăng tải, chỉnh sửa hoặc xóa bài viết tin tức showroom và đánh giá xe."],
                ["Tiền điều kiện", "Người dùng đã đăng nhập với vai trò Admin hoặc Nhân viên được cấp quyền viết bài."],
                ["Hậu điều kiện", "Bài viết được lưu vào bảng posts, revalidate các trang tin tức tĩnh."],
                ["Luồng sự kiện chính", "1. Người đăng bài mở trang quản lý bài viết, click 'Thêm bài viết'.\n2. Nhập tiêu đề, nội dung (hỗ trợ song ngữ), chọn danh mục và tải ảnh đại diện lên.\n3. Chọn trạng thái (nháp hoặc xuất bản) và bấm 'Lưu'.\n4. Server upload ảnh thumbnail lên thư mục lưu trữ cục bộ.\n5. INSERT hoặc UPDATE bản ghi trong bảng posts.\n6. Gọi revalidatePath để cập nhật cache và báo thành công."],
                ["Luồng rẽ nhánh", "4a. Nếu tải ảnh lên thất bại: Báo lỗi và chặn không cho lưu bài viết."]
            ]
        },
        {
            "name": "Xem thống kê báo cáo",
            "seq_img": "seq_analytics.png", "class_img": "class_analytics.png",
            "desc": [
                ["Tên ca sử dụng", "Xem thống kê báo cáo"],
                ["Tác nhân", "Quản trị viên"],
                ["Mô tả vắn tắt", "Tổng hợp doanh số bán hàng, tồn kho và hiệu suất chăm sóc khách hàng toàn showroom."],
                ["Tiền điều kiện", "Đăng nhập bằng tài khoản Quản trị viên (Admin)."],
                ["Hậu điều kiện", "Dữ liệu thống kê được hiển thị trực quan thông qua biểu đồ Recharts."],
                ["Luồng sự kiện chính", "1. Admin truy cập trang thống kê, chọn khoảng thời gian cần báo cáo.\n2. Server tính toán tổng doanh số bằng cách SUM giá bán thực tế của các xe đã bán (status = sold).\n3. Đếm số lượng tồn kho theo hãng xe và đếm số lead CRM thành công.\n4. Trả về cấu trúc dữ liệu thống kê JSON.\n5. Client vẽ biểu đồ cột doanh thu và biểu đồ tròn cơ cấu thương hiệu xe."],
                ["Luồng rẽ nhánh", "2a. Nếu không có dữ liệu bán hàng trong khoảng thời gian đã chọn: Hiển thị biểu đồ trống và các số liệu bằng 0."]
            ]
        },
        {
            "name": "Quản lý tài khoản",
            "seq_img": "seq_block_user.png", "class_img": "class_block_user.png",
            "desc": [
                ["Tên ca sử dụng", "Quản lý tài khoản"],
                ["Tác nhân", "Quản trị viên"],
                ["Mô tả vắn tắt", "Cho phép Admin quản lý danh sách tài khoản, khóa hoặc mở khóa tài khoản người dùng và cập nhật vai trò."],
                ["Tiền điều kiện", "Đăng nhập bằng tài khoản Admin."],
                ["Hậu điều kiện", "Trạng thái hoạt động hoặc vai trò của tài khoản được cập nhật trong database."],
                ["Luồng sự kiện chính", "1. Admin truy cập trang quản lý người dùng.\n2. Hệ thống hiển thị danh sách tài khoản hiện tại từ bảng users.\n3. Admin chọn một tài khoản cụ thể, click 'Khóa tài khoản' (block).\n4. Server UPDATE users SET status = 'blocked' WHERE id.\n5. Đăng xuất cưỡng bức nếu người dùng bị khóa đang online.\n6. Cập nhật lại danh sách trên màn hình quản lý."],
                ["Luồng rẽ nhánh", "3a. Nếu Admin tự khóa tài khoản của chính mình: Hệ thống ngăn chặn và hiển thị thông báo lỗi."]
            ]
        },
        {
            "name": "Quản lý xe",
            "seq_img": "seq_manage_car.png", "class_img": "class_manage_car.png",
            "desc": [
                ["Tên ca sử dụng", "Quản lý xe"],
                ["Tác nhân", "Quản trị viên, Nhân viên"],
                ["Mô tả vắn tắt", "Thêm mới xe, cập nhật thông số kỹ thuật, hình ảnh và trạng thái bán của xe."],
                ["Tiền điều kiện", "Người dùng đã đăng nhập với vai trò Admin hoặc Nhân viên."],
                ["Hậu điều kiện", "Thông tin xe được lưu trong cars và car_images, cập nhật cache trang chủ."],
                ["Luồng sự kiện chính", "1. Người dùng mở trang quản lý xe, click 'Thêm xe'.\n2. Nhập thông tin chi tiết xe (hãng, giá, năm...) và chọn ảnh tải lên.\n3. Bấm 'Lưu'.\n4. Server lưu thông tin vào bảng cars, upload ảnh và lưu URLs vào car_images.\n5. Nếu xe chuyển trạng thái sang 'đã bán' (sold), cập nhật giá bán thực tế và ghi audit log.\n6. Revalidate cache và hiển thị thông báo thành công."],
                ["Luồng rẽ nhánh", "4a. Nếu tải ảnh lên thất bại hoặc thiếu thông số bắt buộc: Hủy bỏ thao tác lưu và báo lỗi."]
            ]
        },
        {
            "name": "So sánh xe",
            "seq_img": "seq_compare_cars.png", "class_img": "class_compare_cars.png",
            "desc": [
                ["Tên ca sử dụng", "So sánh xe"],
                ["Tác nhân", "Khách hàng"],
                ["Mô tả vắn tắt", "Cho phép khách hàng lựa chọn hai xe ô tô khác nhau để so sánh trực quan các thông số kỹ thuật và giá bán."],
                ["Tiền điều kiện", "Khách hàng đang ở trang danh sách xe hoặc chi tiết xe."],
                ["Hậu điều kiện", "Hiển thị bảng so sánh đối chiếu chi tiết thông số của hai xe được chọn."],
                ["Luồng sự kiện chính", "1. Khách hàng click nút 'So sánh' trên tối đa hai chiếc xe mong muốn.\n2. Hệ thống hiển thị thanh trạng thái so sánh và lưu trữ danh sách xe tạm thời.\n3. Khách hàng bấm nút 'So sánh ngay'.\n4. Hệ thống gửi mã ID của hai xe lên server để truy vấn dữ liệu từ bảng cars.\n5. Server truy vấn thông tin chi tiết và trả về đối tượng so sánh.\n6. Hệ thống phân tích, đối chiếu và render bảng so sánh thông số (hãng, giá, động cơ, hộp số...)."],
                ["Luồng rẽ nhánh", "4a. Nếu thông tin một trong hai xe không còn tồn tại (ví dụ đã bị ẩn hoặc xóa): Hệ thống hiển thị thông báo lỗi và quay lại trang danh sách xe."]
            ]
        },
        {
            "name": "Tính toán trả góp",
            "seq_img": "seq_loan_simulation.png", "class_img": "class_loan_simulation.png",
            "desc": [
                ["Tên ca sử dụng", "Tính toán trả góp"],
                ["Tác nhân", "Khách hàng"],
                ["Mô tả vắn tắt", "Hỗ trợ khách hàng mô phỏng kế hoạch tài chính trả góp khi mua xe dựa trên giá xe, tỷ lệ trả trước, lãi suất và kỳ hạn."],
                ["Tiền điều kiện", "Khách hàng đang ở trang chi tiết của xe."],
                ["Hậu điều kiện", "Hiển thị lịch trả nợ chi tiết (gốc, lãi, tổng trả hàng tháng) theo từng kỳ hạn."],
                ["Luồng sự kiện chính", "1. Khách hàng click nút 'Tính trả góp' trên giao diện chi tiết xe.\n2. Hệ thống lấy thông tin giá xe gốc từ bảng cars và hiển thị biểu mẫu tính toán.\n3. Khách hàng nhập các tham số: số tiền trả trước (%), thời gian vay (tháng) và lãi suất năm (%).\n4. Khách hàng bấm nút 'Tính toán'.\n5. Hệ thống gửi yêu cầu tính toán kèm tham số lên server.\n6. Server thực hiện tính toán số gốc, lãi phải trả theo dư nợ giảm dần.\n7. Trả về và hiển thị bảng lịch trả nợ chi tiết hàng tháng."],
                ["Luồng rẽ nhánh", "3a. Nếu khách hàng nhập số tiền trả trước lớn hơn giá trị xe hoặc nhập lãi suất âm: Hệ thống hiển thị thông báo lỗi và yêu cầu nhập lại."]
            ]
        },
        {
            "name": "Tư vấn qua Chat AI",
            "seq_img": "seq_ai_chatbot.png", "class_img": "class_ai_chatbot.png",
            "desc": [
                ["Tên ca sử dụng", "Tư vấn qua Chat AI"],
                ["Tác nhân", "Khách hàng"],
                ["Mô tả vắn tắt", "Cung cấp tính năng trợ lý ảo thông minh sử dụng Gemini AI để giải đáp thắc mắc về xe và thủ tục mua bán của showroom."],
                ["Tiền điều kiện", "Khách hàng mở cửa sổ chat AI trên giao diện website."],
                ["Hậu điều kiện", "Hệ thống hiển thị câu trả lời từ trợ lý ảo AI và lưu lịch sử cuộc trò chuyện vào cơ sở dữ liệu."],
                ["Luồng sự kiện chính", "1. Khách hàng click vào biểu tượng trợ lý ảo để mở Chat AI.\n2. Hệ thống hiển thị giao diện hộp thoại chat.\n3. Khách hàng nhập câu hỏi cần tư vấn (ví dụ: 'Showroom có xe SUV nào dưới 1 tỷ không?') và gửi.\n4. Hệ thống nhận tin nhắn, truy vấn danh sách xe hiện có để làm ngữ cảnh và gửi yêu cầu tới Gemini API.\n5. Gemini API xử lý dữ liệu và trả về nội dung câu trả lời phù hợp.\n6. Hệ thống lưu tin nhắn của khách hàng và phản hồi của AI vào bảng chat_messages.\n7. Hiển thị câu trả lời của AI lên giao diện chat."],
                ["Luồng rẽ nhánh", "4a. Nếu kết nối tới Gemini API bị lỗi hoặc quá hạn (timeout): Hệ thống hiển thị tin nhắn tự động 'Xin lỗi, trợ lý AI đang bận, vui lòng thử lại sau!' và lưu ghi nhận lỗi."]
            ]
        }
    ]

    # LỜI CẢM ƠN
    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_para(doc, "Sau thời gian học tập và rèn luyện tại Trường Đại học Công nghiệp Hà Nội, em đã có cơ hội vận dụng các kiến thức nền tảng về lập trình web, cơ sở dữ liệu, phân tích thiết kế hệ thống và kiểm thử phần mềm vào một bài toán thực tế. Đồ án tốt nghiệp là cột mốc quan trọng giúp em tổng hợp kiến thức, rèn luyện tư duy kỹ thuật và hình thành phương pháp làm việc có hệ thống hơn.", True)
    add_para(doc, "Em xin gửi lời cảm ơn chân thành tới các thầy cô đã giảng dạy, định hướng và hỗ trợ em trong quá trình học tập. Đặc biệt, em xin bày lòng biết ơn sâu sắc tới thầy TS. Phạm Văn Hiệp, người đã trực tiếp hướng dẫn, góp ý và giúp em hoàn thiện đề tài. Những chỉ dẫn của thầy là cơ sở quan trọng để em hoàn thành hệ thống website bán ô tô trực tuyến dành cho showroom TQ Auto.", True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Người thực hiện\nLê Đình Quốc")
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    doc.add_page_break()

    # MỤC LỤC
    add_heading(doc, "MỤC LỤC", 1)
    add_static_toc(doc)
    doc.add_page_break()

    # DANH MỤC TỪ VIẾT TẮT
    add_heading(doc, "DANH MỤC CÁC TỪ VIẾT TẮT", 1)
    add_table(doc, ["STT", "Từ viết tắt", "Ý nghĩa"], [
        [1, "CNTT", "Công nghệ thông tin"],
        [2, "TMĐT", "Thương mại điện tử"],
        [3, "UI", "User Interface - giao diện người dùng"],
        [4, "UX", "User Experience - trải nghiệm người dùng"],
        [5, "API", "Application Programming Interface - giao diện lập trình ứng dụng"],
        [6, "JWT", "JSON Web Token - cơ chế xác thực phiên đăng nhập"],
        [7, "DBMS", "Database Management System - hệ quản trị cơ sở dữ liệu"],
        [8, "CRM", "Customer Relationship Management - quản lý quan hệ khách hàng"],
        [9, "ERD", "Entity Relationship Diagram - sơ đồ thực thể liên kết"],
        [10, "UML", "Unified Modeling Language - ngôn ngữ mô hình hóa thống nhất"],
    ], [900, 2200, 6260])
    doc.add_page_break()

    # DANH MỤC BẢNG BIỂU
    add_heading(doc, "DANH MỤC BẢNG BIỂU", 1)
    table_rows = [
        ("Bảng 1", "Danh mục các từ viết tắt", "4"),
        ("Bảng 1.1", "Công nghệ sử dụng trong hệ thống", "9"),
        ("Bảng 1.2", "Cấu trúc thư mục và chức năng", "10"),
        ("Bảng 1.3", "Các bảng dữ liệu chính", "11"),
        ("Bảng 2.1", "Tác nhân của hệ thống TQ Auto", "13"),
        ("Bảng 2.2", "Bảng tổng hợp yêu cầu chức năng", "14"),
        ("Bảng 2.3", "Bảng tổng hợp yêu cầu phi chức năng", "15"),
    ]
    
    # Add Use Case Spec tables to Table of Tables
    for idx, uc in enumerate(uc_specs):
        table_rows.append((f"Bảng 2.{4+idx}", f"Đặc tả Use Case {uc['name']}", "18"))
    
    # Add DB Tables to Table of Tables
    db_tbls = [
        "users", "cars", "car_images", "appointments", "chat_messages", 
        "posts", "user_favorites", "car_reviews", "customers", "customer_notes", 
        "contact_requests", "notifications"
    ]
    start_tbl_idx = 4 + len(uc_specs)
    for idx, tbl in enumerate(db_tbls):
        table_rows.append((f"Bảng 2.{start_tbl_idx+idx}", f"Thiết kế bảng dữ liệu {tbl}", "29"))
        
    table_rows.extend([
        ("Bảng 3.1", "Môi trường và phiên bản cài đặt", "35"),
        ("Bảng 3.2", "Lịch trình kiểm thử", "38"),
        ("Bảng 3.3", "Yêu cầu tài nguyên phần cứng kiểm thử", "38"),
        ("Bảng 3.4", "Yêu cầu phần mềm phục vụ kiểm thử", "38"),
    ])
    
    # Add Testcases to Table of Tables
    testcases_names = [
        "Đăng ký tài khoản", "Đăng nhập", "Tìm kiếm và Lọc xe", "Đặt lịch xem xe", "Gửi yêu cầu tư vấn",
        "Thêm mới xe (Admin)", "Cập nhật trạng thái lịch hẹn", "Xem thống kê báo cáo", "Quản lý khách hàng CRM", "Chat trực tuyến"
    ]
    for idx, tc_name in enumerate(testcases_names):
        table_rows.append((f"Bảng 3.{5+idx}", f"Ca kiểm thử chức năng {tc_name}", "39"))

    add_table(doc, ["Số hiệu", "Tên bảng", "Trang"], table_rows, [1300, 6560, 1500])
    doc.add_page_break()

    # DANH MỤC HÌNH ẢNH
    add_heading(doc, "DANH MỤC HÌNH ẢNH", 1)
    figure_rows = [
        ("Hình 2.1", "Biểu đồ Use Case tổng quát hệ thống TQ Auto", "16"),
    ]
    # Usecase sub-diagrams
    uc_subs = [
        "Đăng nhập", "Tra cứu & Lọc xe", "Đặt lịch xem xe", "Quản lý xe", "Quản lý khách hàng CRM",
        "Quản lý Lead tư vấn", "Chat trực tuyến", "Thống kê báo cáo", "Quản lý tin tức", "Hồ sơ cá nhân"
    ]
    for idx, uc_sub in enumerate(uc_subs):
        figure_rows.append((f"Hình 2.{2+idx}", f"Biểu đồ Use Case phân hệ {uc_sub}", "16"))
        
    # Seq and Class diagrams
    curr_fig = 12
    for uc in uc_specs:
        figure_rows.append((f"Hình 2.{curr_fig}", f"Biểu đồ tuần tự chức năng {uc['name']}", "18"))
        figure_rows.append((f"Hình 2.{curr_fig+1}", f"Biểu đồ lớp phân tích chức năng {uc['name']}", "18"))
        curr_fig += 2
        
    figure_rows.append((f"Hình 2.{curr_fig}", "Sơ đồ thực thể liên kết (ERD) cơ sở dữ liệu", "29"))
    figure_rows.append((f"Hình 3.1", "Sơ đồ điều hướng giữa các màn hình", "35"))
    
    # Screenshots
    figure_rows.extend([
        ("Hình 3.2", "Giao diện trang chủ TQ Auto", "36"),
        ("Hình 3.3", "Giao diện danh sách xe", "36"),
        ("Hình 3.4", "Chức năng tìm kiếm và lọc xe", "37"),
        ("Hình 3.5", "Giao diện chi tiết xe", "37"),
        ("Hình 3.6", "Dashboard quản trị", "37"),
        ("Hình 3.7", "Khu vực bảng dữ liệu quản trị", "38"),
    ])
    add_table(doc, ["Số hiệu", "Tên hình", "Trang"], figure_rows, [1300, 6560, 1500])
    doc.add_page_break()

    # MỞ ĐẦU
    add_heading(doc, "MỞ ĐẦU", 1)
    add_heading(doc, "1. Lý do chọn đề tài", 2)
    add_para(doc, "Trong bối cảnh chuyển đổi số đang diễn ra mạnh mẽ, website không chỉ là kênh giới thiệu thông tin mà còn trở thành công cụ vận hành, bán hàng và chăm sóc khách hàng của doanh nghiệp. Đối với lĩnh vực kinh doanh ô tô, khách hàng thường cần nhiều thông tin trước khi ra quyết định như giá bán, năm sản xuất, số kilomet đã đi, nguồn gốc xe, tình trạng xe, hình ảnh thực tế, lịch xem xe và tư vấn tài chính. Nếu các dữ liệu này được quản lý rời rạc trên mạng xã hội hoặc bằng bảng tính thủ công, showroom sẽ khó bảo đảm tính chính xác và khó mở rộng quy trình chăm sóc khách hàng.", True)
    add_para(doc, "Từ nhu cầu thực tế đó, đề tài \"Xây dựng website bán ô tô trực tuyến dành cho showroom TQ Auto\" được lựa chọn nhằm xây dựng một hệ thống có khả năng hiển thị kho xe, hỗ trợ khách hàng tìm kiếm và đặt lịch xem xe, đồng thời cung cấp khu vực quản trị để nhân viên và quản trị viên quản lý dữ liệu kinh doanh một cách tập trung.", True)

    add_heading(doc, "2. Mục tiêu của đề tài", 2)
    add_bullets(doc, [
        "Xây dựng website giới thiệu và bán ô tô trực tuyến cho showroom TQ Auto.",
        "Thiết kế giao diện thân thiện, hỗ trợ tiếng Việt và tiếng Anh, hiển thị tốt trên máy tính và thiết bị di động.",
        "Xây dựng các chức năng quản lý xe, khách hàng, lịch hẹn, yêu cầu tư vấn, bài viết tin tức và thông báo.",
        "Hỗ trợ phân quyền người dùng gồm khách hàng, nhân viên và quản trị viên.",
        "Ứng dụng Next.js, React, TypeScript và MySQL để bảo đảm tính hiện đại, dễ bảo trì và có khả năng mở rộng.",
    ])

    add_heading(doc, "3. Đối tượng và phạm vi nghiên cứu", 2)
    add_para(doc, "Đối tượng nghiên cứu của đề tài là quy trình vận hành website bán ô tô trực tuyến, bao gồm nghiệp vụ giới thiệu xe, tiếp nhận yêu cầu tư vấn, đặt lịch xem xe, quản lý khách hàng tiềm năng, quản lý nhân viên và theo dõi thống kê bán hàng. Phân hệ khách hàng hỗ trợ tìm kiếm, xem chi tiết xe và đăng ký lịch xem xe. Phân hệ quản trị và nhân viên hỗ trợ quy trình CRM và quản lý kho xe.", True)

    add_heading(doc, "4. Ý nghĩa khoa học và thực tiễn", 2)
    add_para(doc, "Về mặt khoa học, đề tài giúp vận dụng các kiến thức phân tích thiết kế hệ thống, lập trình giao diện, xử lý phía máy chủ, thiết kế cơ sở dữ liệu quan hệ, bảo mật phiên đăng nhập và kiểm thử chức năng. Về mặt thực tiễn, hệ thống giúp showroom TQ Auto tạo một kênh bán hàng chính thức, nâng cao hình ảnh thương hiệu, giảm thao tác thủ công và tăng khả năng chăm sóc khách hàng theo từng giai đoạn.", True)

    add_heading(doc, "5. Bố cục của báo cáo", 2)
    add_numbered(doc, [
        "Chương 1 trình bày cơ sở lý thuyết và công nghệ sử dụng.",
        "Chương 2 khảo sát, phân tích yêu cầu và thiết kế hệ thống.",
        "Chương 3 trình bày kết quả cài đặt, giao diện và kiểm thử.",
        "Phần kết luận đánh giá kết quả đạt được, hạn chế và hướng phát triển.",
    ])
    doc.add_page_break()

    # CHƯƠNG 1. CƠ SỞ LÝ THUYẾT
    add_heading(doc, "CHƯƠNG 1. CƠ SỞ LÝ THUYẾT", 1)
    add_heading(doc, "1.1. Tổng quan về thương mại điện tử trong lĩnh vực ô tô", 2)
    add_para(doc, "Thương mại điện tử là hình thức tổ chức hoạt động kinh doanh thông qua nền tảng số, cho phép doanh nghiệp giới thiệu sản phẩm, tiếp nhận nhu cầu, xử lý dữ liệu khách hàng và duy trì hoạt động bán hàng không bị giới hạn bởi không gian địa lý. Trong ngành ô tô, hành trình mua xe thường bắt đầu từ việc khách hàng tìm kiếm thông tin trực tuyến, so sánh mẫu xe, tham khảo giá, xem hình ảnh và đánh giá mức độ phù hợp trước khi đến showroom.", True)
    add_para(doc, "Với đặc thù sản phẩm có giá trị cao, website ô tô cần cung cấp thông tin minh bạch và có cấu trúc. Những yếu tố như bộ lọc xe, ảnh chất lượng cao, thông số kỹ thuật, tình trạng còn hàng, nút đặt lịch xem xe và kênh tư vấn trực tiếp ảnh hưởng lớn tới niềm tin của khách hàng. Vì vậy, hệ thống không chỉ là trang giới thiệu mà còn là một công cụ hỗ trợ bán hàng và quản lý quan hệ khách hàng.", True)

    add_heading(doc, "1.2. Công nghệ phát triển ứng dụng", 2)
    add_table(doc, ["Công nghệ", "Vai trò trong hệ thống"], [
        ["Next.js 15", "Framework React hỗ trợ App Router, server components, server actions và tối ưu hóa ứng dụng web."],
        ["React 19", "Xây dựng giao diện dạng component, quản lý trạng thái tương tác và tái sử dụng khối giao diện."],
        ["TypeScript", "Tăng độ tin cậy của mã nguồn nhờ kiểu dữ liệu tĩnh và kiểm tra lỗi sớm."],
        ["Tailwind CSS 4", "Thiết kế giao diện nhanh, nhất quán và đáp ứng nhiều kích thước màn hình."],
        ["MySQL", "Lưu trữ dữ liệu người dùng, xe, ảnh xe, lịch hẹn, khách hàng, tin tức, tin nhắn và thông báo."],
        ["JWT", "Bảo vệ phiên đăng nhập bằng token lưu trong cookie HTTP-only."],
        ["next-intl", "Hỗ trợ đa ngôn ngữ tiếng Việt và tiếng Anh cho khu vực công khai."],
        ["Recharts", "Hiển thị biểu đồ doanh thu, tồn kho, hiệu suất nhân viên và thống kê thương hiệu."],
    ], [2100, 7260])

    add_heading(doc, "1.3. Kiến trúc ứng dụng Next.js", 2)
    add_para(doc, "Dự án được tổ chức theo mô hình ứng dụng Next.js hiện đại, trong đó thư mục src/app chứa các tuyến trang công khai, quản trị và API; thư mục src/components chứa các thành phần giao diện dùng lại; thư mục src/lib chứa lớp xử lý dữ liệu, xác thực, mailer, JWT, hằng số nghiệp vụ và các server actions. Cách tổ chức này giúp tách biệt phần hiển thị, nghiệp vụ và truy cập dữ liệu, giảm trùng lặp và thuận tiện khi mở rộng.", True)
    add_table(doc, ["Nhóm thư mục/tệp", "Chức năng"], [
        ["src/app/[locale]", "Các trang công khai hỗ trợ đa ngôn ngữ: trang chủ, danh sách xe, chi tiết xe, tin tức, liên hệ, đăng nhập, hồ sơ."],
        ["src/app/admin", "Khu vực quản trị dành cho admin: dashboard, xe, khách hàng, lịch hẹn, lead, chat, thống kê, người dùng."],
        ["src/app/staff", "Khu vực nhân viên với quyền hạn hạn chế hơn admin."],
        ["src/components/public", "Header, footer, card xe, bộ lọc, form liên hệ, form đặt lịch, trang chi tiết xe."],
        ["src/components/admin", "Bảng dữ liệu, form quản lý xe, dashboard, biểu đồ, danh sách lead/khách hàng/lịch hẹn."],
        ["src/lib/actions", "Server actions xử lý đăng nhập, xe, khách hàng, lead, lịch hẹn, thống kê, bài viết, thông báo."],
        ["src/lib/db.ts", "Khởi tạo MySQL, định nghĩa bảng, seed dữ liệu và cung cấp hàm CRUD."],
    ], [2500, 6860])

    add_heading(doc, "1.4. Cơ sở dữ liệu MySQL", 2)
    add_para(doc, "MySQL được sử dụng làm hệ quản trị cơ sở dữ liệu quan hệ cho toàn bộ hệ thống. Các bảng chính được tạo tự động khi ứng dụng khởi động gồm users, cars, car_images, appointments, chat_messages, posts, user_favorites, car_reviews, customers, customer_notes, contact_requests và notifications. Cơ sở dữ liệu sử dụng khóa chính dạng UUID, ràng buộc khóa ngoại và các trường trạng thái dạng ENUM để bảo đảm tính nhất quán.", True)
    add_table(doc, ["Bảng", "Dữ liệu lưu trữ chính"], [
        ["users", "Tài khoản, mật khẩu băm, vai trò admin/staff/customer và trạng thái hoạt động."],
        ["cars", "Thông tin xe: hãng, mẫu, năm, giá, số km, nhiên liệu, hộp số, kiểu dáng, màu, trạng thái, ngày giữ chỗ."],
        ["car_images", "Ảnh đại diện và thư viện ảnh của từng xe theo thứ tự hiển thị."],
        ["appointments", "Lịch hẹn xem xe, thông tin khách hàng, thời gian hẹn và trạng thái xử lý."],
        ["customers", "Hồ sơ CRM, xe quan tâm, ngân sách, giai đoạn chăm sóc, nhân viên phụ trách."],
        ["contact_requests", "Yêu cầu tư vấn gửi từ website, loại tư vấn, nhân viên được gán và trạng thái lead."],
        ["posts", "Bài viết tin tức, nội dung song ngữ, thumbnail, danh mục và trạng thái xuất bản."],
        ["notifications", "Thông báo trong hệ thống cho admin, nhân viên hoặc thông báo toàn cục."],
    ], [2000, 7360])

    add_heading(doc, "1.5. Bảo mật và phân quyền", 2)
    add_para(doc, "Hệ thống xác thực người dùng bằng email, mật khẩu đã băm và JWT. Token phiên đăng nhập được lưu trong cookie HTTP-only nhằm hạn chế truy cập từ mã JavaScript phía trình duyệt. Middleware kiểm tra token và vai trò người dùng trước khi cho phép truy cập các tuyến admin, dashboard, staff và profile. Admin có quyền quản lý toàn bộ dữ liệu, nhân viên được truy cập khu vực xử lý khách hàng và lịch hẹn, còn khách hàng sử dụng các chức năng công khai và hồ sơ cá nhân.", True)
    doc.add_page_break()

    # CHƯƠNG 2
    add_heading(doc, "CHƯƠNG 2. KHẢO SÁT, PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    
    add_heading(doc, "2.1. Khảo sát hiện trạng", 2)
    add_para(doc, "Hoạt động kinh doanh tại showroom TQ Auto hiện tại vẫn phụ thuộc lớn vào các phương thức thủ công và tiếp cận truyền thống. Khi có xe mới về showroom, nhân viên phải chụp ảnh và ghi lại thông số kỹ thuật vào sổ tay, hoặc tự đăng lên các trang mạng xã hội cá nhân một cách rời rạc. Điều này khiến cho việc cập nhật thông tin tồn kho không đồng bộ. Có nhiều trường hợp một chiếc xe đã được bán bởi nhân viên này nhưng nhân viên khác vẫn tiếp tục tư vấn cho khách hàng mới, gây ảnh hưởng nghiêm trọng đến uy tín thương hiệu.", True)
    add_para(doc, "Khách hàng muốn tham khảo xe phải trực tiếp đến cửa hàng hoặc gọi điện thoại nhờ nhân viên chụp ảnh gửi qua các ứng dụng tin nhắn. Việc đặt lịch hẹn xem xe hay đăng ký chạy thử cũng được ghi chép thủ công trên sổ sách vật lý, dẫn đến việc dễ bị trùng lặp lịch, bỏ quên lịch hẹn của khách hàng trong giờ cao điểm. Để nâng cao năng lực cạnh tranh và tối ưu hóa quy trình, showroom cần xây dựng một website bán ô tô trực tuyến đồng bộ dữ liệu tập trung, cho phép khách hàng chủ động tra cứu xe, đặt lịch xem xe trực tuyến và nhân viên dễ dàng quản lý thông tin khách hàng (CRM) trên hệ thống.", True)

    add_heading(doc, "2.2. Tác nhân của hệ thống", 2)
    add_para(doc, "Dựa trên các hoạt động thực tế tại showroom, hệ thống TQ Auto được phân quyền chi tiết cho ba tác nhân chính tham gia tương tác trực tiếp bao gồm:", True)
    add_table(doc, ["Tác nhân", "Vai trò trong hệ thống", "Chức năng chính tương tác"], [
        ["Khách hàng", "Người truy cập website để tìm kiếm xe và đặt lịch.", "Xem danh sách xe, tìm kiếm & lọc, xem chi tiết xe, đăng ký, đăng nhập, gửi lịch hẹn, gửi yêu cầu tư vấn, chat trực tuyến."],
        ["Nhân viên", "Nhân viên kinh doanh và chăm sóc khách hàng.", "Xem lịch hẹn được phân công, xử lý lịch hẹn, quản lý hồ sơ CRM khách hàng, thêm ghi chú chăm sóc, chat với khách."],
        ["Quản trị viên", "Admin kiểm soát toàn bộ hệ thống showroom.", "Toàn quyền quản trị: quản lý xe, quản lý tài khoản người dùng, phân công nhân viên, quản lý lịch hẹn, xem thống kê báo cáo doanh số."],
    ], [1500, 2500, 5360])

    add_heading(doc, "2.3. Yêu cầu chức năng", 2)
    add_para(doc, "Các yêu cầu chức năng của hệ thống được chuẩn hóa mã hiệu để thuận tiện cho quá trình phân tích thiết kế và kiểm thử phần mềm:", True)
    add_table(doc, ["Mã chức năng", "Tên chức năng", "Mô tả chi tiết"], [
        ["F01", "Xem danh sách xe", "Hiển thị kho xe đang bán kèm hình ảnh đại diện, giá bán, năm sản xuất, số km, hộp số."],
        ["F02", "Tìm kiếm và lọc xe", "Hỗ trợ khách hàng tìm kiếm xe theo từ khóa và lọc nâng cao theo hãng, khoảng giá, nhiên liệu, tình trạng."],
        ["F03", "Xem chi tiết xe", "Hiển thị thư viện ảnh chi tiết, thông số kỹ thuật đầy đủ, bài viết mô tả chi tiết và danh sách xe liên quan."],
        ["F04", "Đăng ký & Đăng nhập", "Hệ thống xác thực tài khoản người dùng, kiểm tra trạng thái hoạt động của tài khoản."],
        ["F05", "Đặt lịch xem xe", "Khách hàng đăng ký ngày giờ đến xem xe, hệ thống tự động kiểm tra thông tin liên hệ."],
        ["F06", "Quản lý xe (Admin)", "Thêm mới xe, tải lên hình ảnh, cập nhật trạng thái xe (Còn hàng, đã giữ chỗ, đã bán, ẩn xe)."],
        ["F07", "Quản lý khách hàng CRM", "Nhân viên theo dõi thông tin khách hàng, cập nhật giai đoạn (Lead mới, đang tư vấn, đã mua, chăm sóc)."],
        ["F08", "Quản lý lead tư vấn", "Tiếp nhận yêu cầu liên hệ, phân phối nhân viên chăm sóc, ghi nhận trạng thái xử lý."],
        ["F09", "Quản lý lịch hẹn", "Xem danh sách lịch hẹn của khách hàng, cập nhật trạng thái (Chờ duyệt, đã xác nhận, đã hoàn thành, đã hủy)."],
        ["F10", "Chat trực tuyến", "Cung cấp giao diện nhắn tin thời gian thực giữa khách hàng và nhân viên hỗ trợ."],
        ["F11", "Thống kê báo cáo", "Biểu diễn trực quan doanh thu theo tháng/năm, số lượng tồn kho theo hãng xe, hiệu suất bán hàng của nhân viên."],
        ["F12", "Tin tức song ngữ", "Cho phép đăng tải và quản lý bài viết tin tức xe hơi với hai phiên bản ngôn ngữ tiếng Việt và tiếng Anh."],
    ], [1300, 2200, 5860])

    add_heading(doc, "2.4. Yêu cầu phi chức năng", 2)
    add_para(doc, "Nhằm đảm bảo hệ thống vận hành ổn định trên môi trường thực tế và mang lại trải nghiệm tốt nhất, hệ thống cần đáp ứng các tiêu chuẩn phi chức năng sau:", True)
    add_table(doc, ["Nhóm yêu cầu", "Nội dung yêu cầu chi tiết"], [
        ["Hiệu năng (Performance)", "Tốc độ tải trang dưới 2 giây. Hình ảnh xe được nén tối ưu. Thời gian phản hồi API Server Actions dưới 500ms."],
        ["Tính dễ sử dụng (Usability)", "Giao diện hiện đại, trực quan, hỗ trợ thiết kế đáp ứng (Responsive) hiển thị tốt trên cả điện thoại di động và máy tính."],
        ["Bảo mật (Security)", "Mật khẩu người dùng được băm bằng thuật toán bảo mật. Token JWT được lưu trong Cookie với cờ HTTP-only để chống tấn công XSS."],
        ["Khả năng mở rộng", "Kiến trúc Next.js App Router phân tách rõ ràng UI Component, Server Actions giúp dễ dàng tích hợp cổng thanh toán đặt cọc."],
        ["Đa ngôn ngữ (Bilingual)", "Khu vực công khai hỗ trợ chuyển đổi linh hoạt ngôn ngữ giữa tiếng Việt và tiếng Anh nhờ cấu hình module next-intl."],
    ], [2200, 7160])

    # 2.5 Sơ đồ Use Case
    add_heading(doc, "2.5. Sơ đồ Use Case", 2)
    add_para(doc, "Sơ đồ Use Case tổng quát biểu diễn các mối quan hệ tương tác giữa ba tác nhân chính (Khách hàng, Nhân viên, Quản trị viên) với các nhóm chức năng lớn của hệ thống TQ Auto. Các biểu đồ được vẽ đồng bộ dưới dạng thang độ xám (grayscale) chuẩn học thuật:", True)
    add_image(doc, "uc_general.png", "Hình 2.1. Biểu đồ Use Case tổng quát hệ thống TQ Auto")
    
    add_para(doc, "Mối tương tác chi tiết của từng phân hệ tác nhân và các ca sử dụng thành phần với mối quan hệ <<include>> (bao gồm) và <<extend>> (mở rộng) được thể hiện cụ thể:", True)
    
    uc_pairs = [
        ("uc_login.png", "Hình 2.2. Biểu đồ Use Case phân hệ Đăng nhập"),
        ("uc_view_car.png", "Hình 2.3. Biểu đồ Use Case phân hệ Tra cứu & Lọc xe"),
        ("uc_appointment.png", "Hình 2.4. Biểu đồ Use Case phân hệ Đặt lịch xem xe"),
        ("uc_manage_car.png", "Hình 2.5. Biểu đồ Use Case phân hệ Quản lý xe"),
        ("uc_crm.png", "Hình 2.6. Biểu đồ Use Case phân hệ Quản lý khách hàng CRM"),
        ("uc_lead.png", "Hình 2.7. Biểu đồ Use Case phân hệ Quản lý Lead tư vấn"),
        ("uc_chat.png", "Hình 2.8. Biểu đồ Use Case phân hệ Chat trực tuyến"),
        ("uc_analytics.png", "Hình 2.9. Biểu đồ Use Case phân hệ Thống kê báo cáo"),
        ("uc_news.png", "Hình 2.10. Biểu đồ Use Case phân hệ Quản lý tin tức"),
        ("uc_profile.png", "Hình 2.11. Biểu đồ Use Case phân hệ Hồ sơ cá nhân")
    ]
    for img_f, cap_f in uc_pairs:
        add_image(doc, img_f, cap_f)

    # 2.6 Đặc tả & Phân tích chi tiết ca sử dụng
    add_heading(doc, "2.6. Đặc tả & Phân tích chi tiết ca sử dụng", 2)
    add_para(doc, "Dưới đây là phần đặc tả chi tiết luồng sự kiện, sơ đồ tuần tự và sơ đồ lớp phân tích cho 19 ca sử dụng cốt lõi của hệ thống, giúp biểu diễn rõ ràng sự giao tiếp giữa các thành phần giao diện (Boundary), bộ điều khiển (Control) và thực thể cơ sở dữ liệu (Entity):", True)

    curr_fig = 12
    for idx, uc in enumerate(uc_specs):
        num = f"2.6.{idx+1}"
        add_heading(doc, f"{num}. Ca sử dụng {uc['name']}", 3)
        
        # Transform the description to the 8-row format
        new_desc = transform_desc(uc["desc"])
        add_usecase_table(doc, new_desc, [2200, 7160])
        
        seq_cap = f"Hình 2.{curr_fig}. Biểu đồ tuần tự chức năng {uc['name']}"
        class_cap = f"Hình 2.{curr_fig+1}. Biểu đồ lớp phân tích chức năng {uc['name']}"
        add_image(doc, uc["seq_img"], seq_cap)
        add_image(doc, uc["class_img"], class_cap)
        curr_fig += 2

    # 2.7 Thiết kế cơ sở dữ liệu chi tiết
    add_heading(doc, "2.7. Thiết kế cơ sở dữ liệu chi tiết", 2)
    add_para(doc, "Sơ đồ quan hệ thực thể (ERD) thể hiện cấu trúc liên kết khóa chính - khóa ngoại giữa 12 bảng cơ sở dữ liệu của hệ thống website TQ Auto, được vẽ theo phong cách grayscale đồng bộ:", True)
    add_image(doc, "erd.png", f"Hình 2.{curr_fig}. Sơ đồ thực thể liên kết (ERD) cơ sở dữ liệu")
    
    # 2.7.1. Các yêu cầu về dữ liệu
    add_heading(doc, "2.7.1. Các yêu cầu về dữ liệu", 3)
    add_para(doc, "Phần này mô tả các yêu cầu về dữ liệu đối với 12 thực thể chính trong cơ sở dữ liệu hệ thống website TQ Auto, bao gồm ý nghĩa lưu trữ dữ liệu và các mối quan hệ tương quan giữa các thực thể:", True)
    
    add_para(doc, "1. Tài khoản người dùng (users): Lưu trữ thông tin tài khoản bao gồm mã tài khoản (UUID), họ tên, email đăng nhập, số điện thoại, mật khẩu mã hóa, ảnh đại diện, vai trò (Quản trị viên, Nhân viên, Khách hàng) và trạng thái hoạt động. Mỗi tài khoản người dùng có thể thực hiện đăng bán hoặc quản lý nhiều xe ô tô (cars), đặt nhiều lịch hẹn xem xe (appointments), thêm nhiều xe yêu thích (user_favorites), viết nhiều đánh giá xe (car_reviews), ghi nhận nhiều ghi chú khách hàng (customer_notes), được phân công nhiều yêu cầu tư vấn (contact_requests), quản lý nhiều hồ sơ khách hàng (customers), và nhận nhiều thông báo (notifications).", True)
    
    add_para(doc, "2. Thông tin xe ô tô (cars): Lưu trữ thông tin chi tiết về từng mẫu xe đang trưng bày tại showroom bao gồm mã xe (UUID), tên xe, thương hiệu, dòng xe, năm sản xuất, giá bán, số km đã đi, loại nhiên liệu, hộp số, kiểu dáng, màu sắc, số chỗ ngồi, dung tích động cơ, mô tả chi tiết, tình trạng (xe mới/đã qua sử dụng), hệ dẫn động, trạng thái xe (Còn bán, Giữ chỗ, Đã bán, Ẩn), giá bán thực tế và ngày giờ giao dịch. Mỗi xe có thể sở hữu nhiều hình ảnh trong thư viện ảnh (car_images), nhận được nhiều lịch hẹn xem xe (appointments), được đưa vào danh sách yêu thích của nhiều người dùng (user_favorites), nhận nhiều phản hồi đánh giá (car_reviews), hoặc là mẫu xe được quan tâm chính trong các hồ sơ khách hàng (customers).", True)
    
    add_para(doc, "3. Thư viện ảnh xe (car_images): Lưu trữ danh sách các hình ảnh đi kèm của từng xe ô tô để hiển thị slide trình chiếu, bao gồm mã ảnh (UUID), mã xe liên kết, đường dẫn tệp ảnh và thứ tự hiển thị (với thứ tự 0 mặc định là ảnh đại diện). Mỗi ảnh bắt buộc thuộc về duy nhất một xe ô tô (cars).", True)
    
    add_para(doc, "4. Lịch hẹn xem xe (appointments): Lưu trữ thông tin các lượt đăng ký lịch hẹn trực tiếp đến showroom của khách hàng bao gồm mã lịch hẹn (UUID), mã xe quan tâm, mã tài khoản khách hàng (nếu có), họ tên, số điện thoại, email, ngày giờ hẹn, ghi chú đi kèm và trạng thái xử lý (Chờ duyệt, Xác nhận, Hủy, Hoàn thành). Mỗi lịch hẹn thuộc về một xe ô tô (cars) và có thể liên kết với một tài khoản người dùng (users).", True)
    
    add_para(doc, "5. Tin nhắn hỗ trợ (chat_messages): Lưu trữ các đoạn hội thoại tư vấn trực tuyến thời gian thực giữa khách truy cập website và nhân viên hỗ trợ thông qua mã phiên chat (session_id), vai trò người gửi (nhân viên/khách hàng), họ tên người gửi, nội dung tin nhắn, trạng thái đã đọc và thời gian gửi. Các tin nhắn có cùng mã phiên chat sẽ thuộc về một cuộc trò chuyện duy nhất và có thể được liên kết với hồ sơ khách hàng CRM tương ứng (customers).", True)
    
    add_para(doc, "6. Tin tức song ngữ (posts): Lưu trữ các bài viết tin tức, đánh giá xe, sự kiện của showroom hỗ trợ hiển thị song ngữ với tiêu đề tiếng Việt/Anh, slug thân thiện cho SEO, nội dung tiếng Việt/Anh, ảnh đại diện, danh mục, số lượt đọc và trạng thái xuất bản (nháp/công khai). Đây là thực thể độc lập phục vụ nhu cầu marketing và cung cấp thông tin cho người dùng website.", True)
    
    add_para(doc, "7. Danh sách xe yêu thích (user_favorites): Lưu giữ mối liên kết giữa tài khoản người dùng và những chiếc xe họ quan tâm, yêu thích để xem lại sau, bao gồm khóa ngoại trỏ tới tài khoản người dùng (users) và xe ô tô (cars). Mỗi bản ghi biểu diễn hành vi lưu xe của một người dùng đối với một chiếc xe cụ thể.", True)
    
    add_para(doc, "8. Đánh giá xe (car_reviews): Lưu trữ các phản hồi, đánh giá chất lượng và bình luận của khách hàng về các mẫu xe ô tô, gồm mã đánh giá (UUID), tài khoản đánh giá, xe được đánh giá, điểm số đánh giá (1-5 sao), nội dung bình luận và thời gian đăng tải. Mỗi đánh giá thuộc về một tài khoản người dùng (users) và một xe ô tô (cars) nhất định.", True)
    
    add_para(doc, "9. Hồ sơ khách hàng CRM (customers): Lưu trữ thông tin hồ sơ khách hàng tiềm năng được quản lý trong hệ thống CRM phục vụ việc chăm sóc sau bán hàng, bao gồm mã hồ sơ (UUID), họ tên, số điện thoại, email, xe đang quan tâm, ngân sách dự kiến, giai đoạn chăm sóc (Lead mới, Đang tư vấn, Đã hẹn lịch, Đã báo giá, Đang thương lượng, Đã đặt cọc, Đã mua, Chăm sóc sau bán), nhân viên được gán phụ trách, nguồn khách hàng và mã phiên chat liên kết. Mỗi hồ sơ được phụ trách bởi một nhân viên (users), liên kết với xe họ quan tâm (cars), và có thể có nhiều ghi chú chăm sóc (customer_notes).", True)
    
    add_para(doc, "10. Ghi chú chăm sóc khách hàng (customer_notes): Lưu trữ lịch sử chăm sóc và ghi nhận ý kiến khách hàng sau mỗi cuộc gọi hoặc trao đổi trực tiếp, bao gồm mã ghi chú (UUID), mã hồ sơ khách hàng liên kết, mã nhân viên thực hiện ghi chú, nội dung chi tiết cuộc trò chuyện và thời gian thực hiện. Mỗi ghi chú thuộc về một khách hàng (customers) và được tạo bởi một nhân viên (users).", True)
    
    add_para(doc, "11. Yêu cầu liên hệ / Lead (contact_requests): Lưu trữ thông tin đăng ký nhận tư vấn từ các biểu mẫu liên hệ nhanh trên website, bao gồm mã yêu cầu (UUID), họ tên, số điện thoại, email, nhu cầu tư vấn (Giá xe, Thủ tục trả góp, Đăng ký lái thử), nội dung tin nhắn, nhân viên kinh doanh tiếp nhận xử lý và trạng thái xử lý của Lead. Mỗi yêu cầu liên hệ có thể được phân công cho một nhân viên kinh doanh (users) phụ trách.", True)
    
    add_para(doc, "12. Thông báo hệ thống (notifications): Lưu trữ các thông báo đẩy tự động khi phát sinh sự kiện (như có lịch hẹn mới, có yêu cầu tư vấn mới), bao gồm mã thông báo (UUID), tài khoản nhận thông báo (NULL nếu thông báo tới toàn showroom), tiêu đề, nội dung chi tiết, liên kết chuyển hướng khi nhấp chuột, trạng thái đã đọc và thời gian gửi. Mỗi thông báo thuộc về một tài khoản người dùng (users) nhận thông tin.", True)

    add_para(doc, "Chi tiết định nghĩa các trường thông tin, kiểu dữ liệu, các ràng buộc khóa của từng bảng được mô tả trong các bảng thiết kế chi tiết dưới đây:", True)

    # 12 DB tables definition tables
    # 1. users
    add_heading(doc, "2.7.2. Bảng lưu thông tin tài khoản (users)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "Mã định danh duy nhất dạng UUID"],
        ["full_name", "VARCHAR(255)", "NOT NULL", "Họ và tên người dùng"],
        ["email", "VARCHAR(255)", "UNIQUE NOT NULL", "Email đăng nhập hệ thống"],
        ["phone", "VARCHAR(50)", "NOT NULL", "Số điện thoại liên lạc"],
        ["password", "VARCHAR(255)", "NOT NULL", "Mật khẩu đã được băm bằng sha256"],
        ["avatar", "TEXT", "NULL", "Đường dẫn ảnh đại diện người dùng"],
        ["role", "ENUM('admin','staff','customer')", "NOT NULL", "Phân quyền (Quản trị, Nhân viên, Khách hàng)"],
        ["status", "ENUM('active','blocked')", "DEFAULT 'active'", "Trạng thái tài khoản (Đang hoạt động, Tạm khóa)"],
        ["created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời gian khởi tạo tài khoản"],
    ], [1800, 2200, 2200, 3160])

    # 2. cars
    add_heading(doc, "2.7.3. Bảng lưu thông tin xe ô tô (cars)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "Mã định danh duy nhất của xe dạng UUID"],
        ["user_id", "VARCHAR(36)", "FK references users(id)", "ID của người đăng xe (thường là Admin)"],
        ["title", "VARCHAR(255)", "NOT NULL", "Tên xe hiển thị mặc định"],
        ["brand", "VARCHAR(100)", "NOT NULL", "Thương hiệu sản xuất xe (Toyota, Ford...)"],
        ["model", "VARCHAR(100)", "NOT NULL", "Dòng xe (Camry, Ranger, Civic...)"],
        ["year", "INT", "NOT NULL", "Năm sản xuất xe"],
        ["price", "BIGINT", "NOT NULL", "Giá bán xe (VND)"],
        ["mileage", "INT", "NOT NULL", "Số kilomet xe đã di chuyển"],
        ["fuel_type", "VARCHAR(50)", "NOT NULL", "Loại nhiên liệu (Xăng, Dầu, Điện)"],
        ["transmission", "VARCHAR(50)", "NOT NULL", "Hộp số (Số sàn, Số tự động)"],
        ["body_type", "VARCHAR(50)", "NOT NULL", "Kiểu dáng xe (Sedan, SUV...)"],
        ["color", "VARCHAR(50)", "NOT NULL", "Màu sắc ngoại thất xe"],
        ["seats", "INT", "NOT NULL", "Số chỗ ngồi đăng ký"],
        ["engine", "VARCHAR(100)", "NOT NULL", "Dung tích động cơ (2.0L, 1.5L...)"],
        ["description", "TEXT", "NOT NULL", "Nội dung bài viết mô tả chi tiết xe"],
        ["status", "ENUM('avail','res','sold','hid')", "DEFAULT 'avail'", "Trạng thái xe (Còn bán, Giữ chỗ, Đã bán, Ẩn)"],
        ["car_condition", "ENUM('new','used')", "DEFAULT 'used'", "Tình trạng xe (Xe mới, Xe lướt / đã qua sử dụng)"],
        ["drivetrain", "ENUM('FWD','RWD','AWD','4WD')", "DEFAULT 'FWD'", "Hệ dẫn động xe"],
        ["sold_price", "BIGINT", "NULL", "Giá bán thực tế (cho analytics)"],
        ["sold_at", "DATETIME", "NULL", "Thời gian ghi nhận bán xe"],
    ], [1800, 2200, 2200, 3160])

    # 3. car_images
    add_heading(doc, "2.7.4. Bảng lưu thư viện ảnh xe (car_images)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "Mã ảnh dạng UUID"],
        ["car_id", "VARCHAR(36)", "FK references cars(id)", "ID xe sở hữu hình ảnh này"],
        ["image_url", "TEXT", "NOT NULL", "Đường dẫn ảnh lưu trữ trên server/cloud"],
        ["sort_order", "INT", "NOT NULL", "Thứ tự sắp xếp hiển thị ảnh (0 là ảnh đại diện)"],
        ["created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời gian tải ảnh lên"],
    ], [1800, 2200, 2200, 3160])

    # 4. appointments
    add_heading(doc, "2.7.5. Bảng lưu lịch hẹn xem xe (appointments)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "Mã lịch hẹn UUID"],
        ["car_id", "VARCHAR(36)", "FK references cars(id)", "ID chiếc xe khách đăng ký xem"],
        ["user_id", "VARCHAR(36)", "FK references users(id) NULL", "ID tài khoản khách hàng đăng nhập (nếu có)"],
        ["customer_name", "VARCHAR(255)", "NOT NULL", "Họ tên khách hàng đăng ký lịch"],
        ["customer_phone", "VARCHAR(50)", "NOT NULL", "Số điện thoại khách hàng"],
        ["customer_email", "VARCHAR(255)", "NOT NULL", "Email nhận thông tin xác nhận"],
        ["appointment_date", "TIMESTAMP", "NOT NULL", "Ngày giờ hẹn đến showroom"],
        ["note", "TEXT", "NULL", "Ghi chú bổ sung từ khách hàng"],
        ["status", "ENUM('pend','conf','canc','comp')", "DEFAULT 'pend'", "Trạng thái (Chờ duyệt, Xác nhận, Hủy, Hoàn thành)"],
        ["created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời gian gửi lịch hẹn"],
    ], [1800, 2200, 2200, 3160])

    # 5. chat_messages
    add_heading(doc, "2.7.6. Bảng lưu tin nhắn hỗ trợ (chat_messages)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "ID tin nhắn dạng UUID"],
        ["session_id", "VARCHAR(50)", "NOT NULL", "Mã phiên chat duy nhất của khách truy cập"],
        ["sender_role", "ENUM('customer','staff')", "NOT NULL", "Vai trò người gửi tin"],
        ["sender_name", "VARCHAR(255)", "NOT NULL", "Tên hiển thị của người gửi"],
        ["message_text", "TEXT", "NOT NULL", "Nội dung tin nhắn gửi đi"],
        ["is_read", "TINYINT(1)", "DEFAULT 0", "Đánh dấu tin nhắn đã đọc (0: Chưa đọc, 1: Đã đọc)"],
        ["created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời gian gửi tin nhắn"],
    ], [1800, 2200, 2200, 3160])

    # 6. posts
    add_heading(doc, "2.7.7. Bảng lưu tin tức song ngữ (posts)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "ID bài viết UUID"],
        ["title_vi", "VARCHAR(255)", "NULL", "Tiêu đề tiếng Việt"],
        ["title_en", "VARCHAR(255)", "NULL", "Tiêu đề tiếng Anh"],
        ["slug", "VARCHAR(255)", "UNIQUE NOT NULL", "Đường dẫn thân thiện của bài viết"],
        ["content_vi", "LONGTEXT", "NULL", "Nội dung bài viết tiếng Việt"],
        ["content_en", "LONGTEXT", "NULL", "Nội dung bài viết tiếng Anh"],
        ["thumbnail", "TEXT", "NOT NULL", "Đường dẫn ảnh đại diện bài viết"],
        ["category", "VARCHAR(100)", "NOT NULL", "Danh mục tin tức (Đánh giá, Sự kiện...)"],
        ["views", "INT", "DEFAULT 0", "Tổng số lượt đọc bài viết"],
        ["published", "TINYINT(1)", "DEFAULT 0", "Trạng thái xuất bản (0: Nháp, 1: Công khai)"],
    ], [1800, 2200, 2200, 3160])

    # 7. user_favorites
    add_heading(doc, "2.7.8. Bảng lưu xe yêu thích (user_favorites)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["user_id", "VARCHAR(36)", "PK, FK users(id)", "ID người dùng yêu thích"],
        ["car_id", "VARCHAR(36)", "PK, FK cars(id)", "ID xe được đưa vào danh sách yêu thích"],
    ], [1800, 2200, 2200, 3160])

    # 8. car_reviews
    add_heading(doc, "2.7.9. Bảng lưu đánh giá xe (car_reviews)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "ID đánh giá UUID"],
        ["user_id", "VARCHAR(36)", "FK references users(id)", "ID khách hàng viết đánh giá"],
        ["car_id", "VARCHAR(36)", "FK references cars(id)", "ID xe được đánh giá"],
        ["rating", "INT", "NOT NULL", "Số sao đánh giá (từ 1 đến 5)"],
        ["comment", "TEXT", "NOT NULL", "Nội dung bình luận đánh giá"],
        ["created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời gian đánh giá"],
    ], [1800, 2200, 2200, 3160])

    # 9. customers
    add_heading(doc, "2.7.10. Bảng lưu hồ sơ khách hàng CRM (customers)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "ID hồ sơ khách hàng UUID"],
        ["full_name", "VARCHAR(255)", "NOT NULL", "Họ và tên khách hàng"],
        ["phone", "VARCHAR(50)", "NOT NULL", "Số điện thoại chính xác"],
        ["email", "VARCHAR(255)", "NOT NULL", "Địa chỉ email"],
        ["interested_car_id", "VARCHAR(36)", "FK references cars(id) NULL", "Chiếc xe khách hàng đang quan tâm nhiều nhất"],
        ["budget", "VARCHAR(100)", "NULL", "Khoảng ngân sách dự kiến của khách"],
        ["stage", "ENUM('new_lead','consulting','appointment','quotation','negotiating','reserved','purchased','follow_up')", "DEFAULT 'new_lead'", "Giai đoạn chăm sóc khách hàng trong phễu CRM"],
        ["assigned_staff_id", "VARCHAR(36)", "FK references users(id) NULL", "Nhân viên được gán phụ trách chăm sóc khách"],
        ["source", "VARCHAR(100)", "DEFAULT 'website'", "Nguồn khách hàng (website, Facebook, giới thiệu)"],
        ["session_id", "VARCHAR(100)", "NULL", "Mã phiên chat kết nối với chat_messages"],
    ], [1800, 2200, 2200, 3160])

    # 10. customer_notes
    add_heading(doc, "2.7.11. Bảng lưu ghi chú chăm sóc (customer_notes)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "ID ghi chú UUID"],
        ["customer_id", "VARCHAR(36)", "FK references customers(id)", "ID hồ sơ khách hàng được ghi chú"],
        ["staff_id", "VARCHAR(36)", "FK references users(id)", "ID nhân viên thực hiện cuộc gọi/trao đổi và ghi chú"],
        ["content", "TEXT", "NOT NULL", "Nội dung chi tiết cuộc trao đổi chăm sóc khách"],
        ["created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời gian ghi chú"],
    ], [1800, 2200, 2200, 3160])

    # 11. contact_requests
    add_heading(doc, "2.7.12. Bảng lưu yêu cầu liên hệ / Lead (contact_requests)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "ID yêu cầu tư vấn UUID"],
        ["full_name", "VARCHAR(255)", "NOT NULL", "Họ và tên khách hàng gửi liên hệ"],
        ["phone", "VARCHAR(50)", "NOT NULL", "Số điện thoại"],
        ["email", "VARCHAR(255)", "NOT NULL", "Email liên hệ"],
        ["consultation_type", "VARCHAR(100)", "NOT NULL", "Loại hình tư vấn (Giá xe, Thủ tục trả góp, Lái thử xe)"],
        ["message", "TEXT", "NOT NULL", "Nội dung câu hỏi yêu cầu hỗ trợ"],
        ["assigned_staff_id", "VARCHAR(36)", "FK references users(id) NULL", "Nhân viên kinh doanh tiếp nhận xử lý"],
        ["stage", "ENUM('new_lead','assigned','consulting','appointment','quotation','purchased','closed')", "DEFAULT 'new_lead'", "Trạng thái xử lý của Lead"],
    ], [1800, 2200, 2200, 3160])

    # 12. notifications
    add_heading(doc, "2.7.13. Bảng lưu thông báo hệ thống (notifications)", 3)
    add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả ý nghĩa"], [
        ["id", "VARCHAR(36)", "PRIMARY KEY", "ID thông báo UUID"],
        ["user_id", "VARCHAR(36)", "FK references users(id) NULL", "ID tài khoản nhận thông báo (NULL nếu báo toàn cục)"],
        ["title", "VARCHAR(255)", "NOT NULL", "Tiêu đề thông báo ngắn gọn"],
        ["content", "TEXT", "NOT NULL", "Chi tiết nội dung thông báo"],
        ["link", "VARCHAR(255)", "NULL", "Đường dẫn trang web liên kết khi click thông báo"],
        ["is_read", "TINYINT(1)", "DEFAULT 0", "Trạng thái đọc (0: Chưa đọc, 1: Đã đọc)"],
        ["created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời gian phát sinh sự kiện thông báo"],
    ], [1800, 2200, 2200, 3160])

    add_heading(doc, "2.8. Thiết kế giao diện tổng quan", 2)
    add_para(doc, "Giao diện hệ thống website TQ Auto được xây dựng theo ngôn ngữ thiết kế tối giản, chuyên nghiệp và lịch lãm nhằm làm nổi bật giá trị của sản phẩm ô tô trưng bày. Tông màu chủ đạo sử dụng là sự kết hợp giữa Trắng (Bản nền chính), Đen/Xám đậm (Hệ thống chữ hiển thị) và Đỏ đậm làm màu nhấn cho các nút kêu gọi hành động quan trọng (như nút Đặt lịch hẹn, Đăng ký tư vấn).", True)
    add_para(doc, "Hệ thống bố cục được tổ chức dạng lưới linh hoạt (CSS Grid & Flexbox) tích hợp thuộc tính responsive tự động thích ứng với màn hình điện thoại di động và máy tính bảng. Các khối chức năng như header (thanh điều hướng), footer, thẻ thông tin xe (Car Card), bộ lọc xe (Filter Sidebar) được viết dưới dạng các React Component độc lập có khả năng tái sử dụng cao ở nhiều trang giao diện khác nhau, giúp giảm thiểu độ trùng lặp mã nguồn và nâng cao tính dễ bảo trì của ứng dụng.", True)
    doc.add_page_break()

    # CHƯƠNG 3
    add_heading(doc, "CHƯƠNG 3. KẾT QUẢ CÀI ĐẶT VÀ KIỂM THỬ HỆ THỐNG", 1)
    
    add_heading(doc, "3.1. Sơ đồ điều hướng màn hình", 2)
    add_para(doc, "Sơ đồ dưới đây thể hiện luồng chuyển đổi trạng thái giao diện và điều hướng của người dùng từ trang chủ đến các phân hệ công khai và trang quản trị dashboard của showroom:", True)
    add_image(doc, "navigation_chart.png", "Hình 3.1. Sơ đồ điều hướng giữa các màn hình")
    add_para(doc, "Khi khách hàng truy cập Trang chủ, họ có thể dẫn đến Trang danh sách xe hoặc gửi yêu cầu tại Trang liên hệ. Từ chi tiết một chiếc xe cụ thể, khách hàng có thể mở biểu mẫu đặt lịch xem xe. Trang đăng nhập hoạt động như một cổng bảo mật, kiểm tra thông tin JWT và phân quyền điều hướng người quản trị vào Dashboard Admin và nhân viên bán hàng vào Dashboard Staff.", True)

    add_heading(doc, "3.2. Thiết kế chi tiết màn hình", 2)
    add_para(doc, "Dưới đây là hình ảnh chụp trực quan các giao diện chính của hệ thống TQ Auto đã được cài đặt và vận hành thực tế:", True)

    add_heading(doc, "3.2.1. Giao diện trang chủ và danh sách xe phân hệ công khai", 3)
    add_para(doc, "Trang chủ TQ Auto được thiết kế sang trọng với banner xe cao cấp, thanh tìm kiếm nhanh hỗ trợ khách hàng định vị phân khúc xe mong muốn ngay lập tức:", True)
    add_image(doc, "home.png", "Hình 3.2. Giao diện trang chủ TQ Auto")
    
    add_para(doc, "Trang danh sách xe hiển thị lưới các xe đang mở bán kèm các bộ lọc thông minh (thương hiệu, khoảng giá, kiểu dáng, nhiên liệu) giúp việc so sánh xe trực quan hơn:", True)
    add_image(doc, "cars.png", "Hình 3.3. Giao diện danh sách xe")
    add_image(doc, "cars-search.png", "Hình 3.4. Chức năng tìm kiếm và lọc xe")

    add_heading(doc, "3.2.2. Giao diện trang chi tiết sản phẩm xe ô tô", 3)
    add_para(doc, "Trang chi tiết cung cấp thư viện ảnh trượt đa góc cạnh, hiển thị bảng thông số kỹ thuật chuẩn quốc tế, bài viết mô tả chi tiết và tích hợp form đặt lịch hẹn nhanh tại góc phải màn hình:", True)
    add_image(doc, "detail-fixed.png", "Hình 3.5. Giao diện chi tiết xe")

    add_heading(doc, "3.2.3. Giao diện hệ thống quản lý Dashboard cho showroom", 3)
    add_para(doc, "Khu vực quản trị hiển thị các thẻ thống kê tổng quan doanh thu, số lượng xe, số lượng lịch hẹn phát sinh trong ngày cùng biểu đồ doanh số trực quan:", True)
    add_image(doc, "admin.png", "Hình 3.6. Dashboard quản trị")
    
    add_para(doc, "Bảng quản lý dữ liệu cung cấp công cụ lọc nâng cao, các nút tác vụ thêm/sửa/xóa và cập nhật trạng thái phễu CRM khách hàng:", True)
    add_image(doc, "admin2.png", "Hình 3.7. Khu vực bảng dữ liệu quản trị")

    # 3.3 Kiểm thử
    add_heading(doc, "3.3. Kiểm thử hệ thống", 2)
    add_para(doc, "Kiểm thử phần mềm được thực hiện nhằm xác minh các chức năng của hệ thống vận hành đúng logic thiết kế và không phát sinh lỗi ngoại lệ khi người dùng nhập sai dữ liệu.", True)
    
    add_heading(doc, "3.3.1. Kế hoạch và môi trường kiểm thử", 3)
    add_para(doc, "Kế hoạch kiểm thử được phân phối chi tiết cho từng giai đoạn phát triển phần mềm bao gồm chuẩn bị tài nguyên và thiết lập môi trường giả lập:", True)
    
    add_table(doc, ["Giai đoạn", "Công việc thực hiện", "Người chịu trách nhiệm"], [
        ["Giai đoạn 1", "Xây dựng các kịch bản kiểm thử (Test Cases) chi tiết cho từng chức năng.", "Lê Đình Quốc"],
        ["Giai đoạn 2", "Kiểm thử hộp đen (Black-box testing) các chức năng đăng nhập, đăng ký, lọc xe.", "Lê Đình Quốc"],
        ["Giai đoạn 3", "Kiểm thử các luồng liên kết: đặt lịch hẹn -> gửi thông báo -> CRM chăm sóc khách hàng.", "Lê Đình Quốc"],
        ["Giai đoạn 4", "Kiểm thử hiệu năng tải trang, kiểm tra lỗi bảo mật định dạng dữ liệu đầu vào SQL.", "Lê Đình Quốc"],
    ])
    
    add_para(doc, "Yêu cầu tài nguyên phần cứng, phần mềm giả lập để tiến hành kiểm thử bao gồm:", True)
    add_table(doc, ["Loại tài nguyên", "Tên thiết bị / Phần mềm cấu hình", "Yêu cầu tối thiểu"], [
        ["Phần cứng", "Máy tính xách tay chạy kiểm thử", "Intel Core i5, 8GB RAM, SSD 256GB"],
        ["Phần mềm", "Hệ điều hành kiểm thử máy client", "Windows 11 / macOS Sequoia"],
        ["Trình duyệt", "Trình duyệt kiểm thử giao diện công khai", "Google Chrome phiên bản mới nhất, Microsoft Edge"],
        ["Database Server", "Hệ quản trị cơ sở dữ liệu backend", "MySQL Server 8.0"],
        ["Công cụ Test API", "Kiểm tra định dạng request Server Actions", "Postman / Thunder Client VS Code"],
    ])

    add_heading(doc, "3.3.2. Kịch bản kiểm thử (Test Cases) chi tiết", 3)
    add_para(doc, "Dưới đây là 10 bảng kịch bản kiểm thử chi tiết mô tả dữ liệu đầu vào và kết quả mong đợi đối với các chức năng cốt lõi của website TQ Auto:", True)

    # TC 1: Đăng ký
    add_heading(doc, "3.3.2.1. Kịch bản kiểm thử chức năng Đăng ký tài khoản", 3)
    add_table(doc, ["Bước thực hiện", "Dữ liệu nhập vào", "Kết quả mong đợi", "Trạng thái thực tế"], [
        ["1. Truy cập trang đăng ký, điền thông tin hợp lệ.", "Họ tên: Nguyễn Văn A\nEmail: khachhangA@gmail.com\nSĐT: 0987654321\nMật khẩu: KhachhangA123", "Hệ thống lưu tài khoản mới vào bảng users, hiển thị thông báo thành công và tự động đăng nhập.", "Đạt (Pass)"],
        ["2. Đăng ký lại với Email đã tồn tại.", "Họ tên: Nguyễn Văn B\nEmail: khachhangA@gmail.com\nSĐT: 0912345678\nMật khẩu: Pass123", "Hệ thống từ chối đăng ký và hiển thị cảnh báo đỏ: 'Email này đã được sử dụng trên hệ thống'.", "Đạt (Pass)"],
    ], [1800, 2200, 5000, 1000])

    # TC 2: Đăng nhập
    add_heading(doc, "3.3.2.2. Kịch bản kiểm thử chức năng Đăng nhập", 3)
    add_table(doc, ["Bước thực hiện", "Dữ liệu nhập vào", "Kết quả mong đợi", "Trạng thái thực tế"], [
        ["1. Nhập đúng Email và Mật khẩu đã đăng ký.", "Email: admin@tqauto.vn\nMật khẩu: admin123", "Đăng nhập thành công, hệ thống sinh JWT lưu trong Cookie HTTP-only, chuyển hướng về Dashboard Admin.", "Đạt (Pass)"],
        ["2. Nhập sai mật khẩu của tài khoản.", "Email: admin@tqauto.vn\nMật khẩu: wrongpass", "Hệ thống báo lỗi: 'Thông tin đăng nhập không chính xác', không sinh token phiên làm việc.", "Đạt (Pass)"],
    ], [1800, 2200, 5000, 1000])

    # TC 3: Tìm kiếm lọc xe
    add_heading(doc, "3.3.2.3. Kịch bản kiểm thử chức năng Tìm kiếm và Lọc xe", 3)
    add_table(doc, ["Bước thực hiện", "Dữ liệu nhập vào", "Kết quả mong đợi", "Trạng thái thực tế"], [
        ["1. Chọn bộ lọc Hãng xe 'Toyota' và khoảng giá dưới 1 tỷ.", "Hãng: Toyota\nMức giá: < 1.000.000.000đ", "Lưới danh sách xe lập tức hiển thị các xe của hãng Toyota có giá trị dưới 1 tỷ, ẩn các dòng xe khác.", "Đạt (Pass)"],
        ["2. Nhập từ khóa không có trong danh sách showroom.", "Từ khóa: 'Lamborghini'", "Hệ thống hiển thị thông báo: 'Không tìm thấy chiếc xe nào phù hợp với yêu cầu của bạn'.", "Đạt (Pass)"],
    ], [1800, 2200, 5000, 1000])

    # TC 4: Đặt lịch xem xe
    add_heading(doc, "3.3.2.4. Kịch bản kiểm thử chức năng Đặt lịch xem xe", 3)
    add_table(doc, ["Bước thực hiện", "Dữ liệu nhập vào", "Kết quả mong đợi", "Trạng thái thực tế"], [
        ["1. Click nút Đặt lịch tại chi tiết xe Camry, điền thông tin đầy đủ.", "Họ tên: Trần Văn C\nSĐT: 0909090909\nNgày hẹn: 15/06/2026\nNote: Xem xe buổi sáng", "Hệ thống lưu lịch hẹn ở trạng thái Chờ duyệt (pending), gửi email xác nhận tự động cho khách.", "Đạt (Pass)"],
        ["2. Gửi lịch hẹn với số điện thoại sai định dạng.", "SĐT: 'abcd1234'", "Hệ thống báo lỗi định dạng số điện thoại, chặn không cho submit form.", "Đạt (Pass)"],
    ], [1800, 2200, 5000, 1000])

    # TC 5: Gửi yêu cầu tư vấn
    add_heading(doc, "3.3.2.5. Kịch bản kiểm thử chức năng Gửi yêu cầu tư vấn (Lead)", 3)
    add_table(doc, ["Bước thực hiện", "Dữ liệu nhập vào", "Kết quả mong đợi", "Trạng thái thực tế"], [
        ["1. Nhập thông tin form liên hệ nhận tư vấn trả góp xe Ranger.", "Họ tên: Lê Văn D\nSĐT: 0911223344\nLoại tư vấn: Thủ tục trả góp\nNội dung: Cần tư vấn lãi suất 80%", "Hệ thống tạo yêu cầu tư vấn mới trong bảng contact_requests, đồng thời tạo một Lead mới trong bảng customers CRM.", "Đạt (Pass)"],
    ], [1800, 2200, 5000, 1000])

    # TC 6: Thêm mới xe (Admin)
    add_heading(doc, "3.3.2.6. Kịch bản kiểm thử chức năng Thêm mới xe (Admin)", 3)
    add_table(doc, ["Bước thực hiện", "Dữ liệu nhập vào", "Kết quả mong đợi", "Trạng thái thực tế"], [
        ["1. Admin mở form thêm xe, nhập đầy đủ thông số hợp lệ.", "Tên: Mercedes C200\nHãng: Mercedes\nNăm: 2024\nGiá: 1.700.000.000đ\nẢnh: c200_front.jpg", "Lưu thông tin thành công vào bảng cars, bảng danh sách xe ngoài trang chủ xuất hiện thêm Mercedes C200 lập tức.", "Đạt (Pass)"],
    ], [1800, 2200, 5000, 1000])

    # TC 7: Cập nhật lịch hẹn
    add_heading(doc, "3.3.2.7. Kịch bản kiểm thử chức năng Cập nhật trạng thái lịch hẹn", 3)
    add_table(doc, ["Bước thực hiện", "Dữ liệu nhập vào", "Kết quả mong đợi", "Trạng thái thực tế"], [
        ["1. Nhân viên mở lịch hẹn của khách Trần Văn C, bấm 'Xác nhận'.", "Hành động: Chọn trạng thái 'Đã xác nhận' và lưu.", "Trường status của appointments chuyển sang 'confirmed', hệ thống gửi email báo lịch hẹn đã duyệt cho khách.", "Đạt (Pass)"],
    ], [1800, 2200, 5000, 1000])

    # TC 8: Xem thống kê
    add_heading(doc, "3.3.2.8. Kịch bản kiểm thử chức năng Xem thống kê báo cáo", 3)
    add_table(doc, ["Bước thực hiện", "Dữ liệu nhập vào", "Kết quả mong đợi", "Trạng thái thực tế"], [
        ["1. Admin chọn khoảng thời gian xem báo cáo doanh số Quý 1/2026.", "Thời gian: 01/01/2026 - 31/03/2026", "Hệ thống tổng hợp đúng doanh thu bán xe, biểu đồ Recharts hiển thị chính xác các cột doanh số tăng trưởng.", "Đạt (Pass)"],
    ], [1800, 2200, 5000, 1000])

    # TC 9: Quản lý khách hàng CRM
    add_heading(doc, "3.3.2.9. Kịch bản kiểm thử chức năng Quản lý khách hàng CRM", 3)
    add_table(doc, ["Bước thực hiện", "Dữ liệu nhập vào", "Kết quả mong đợi", "Trạng thái thực tế"], [
        ["1. Nhân viên chuyển phễu khách hàng Lê Văn D sang 'Đang tư vấn' và nhập ghi chú mới.", "Trạng thái: consulting\nGhi chú: Đã gọi điện tư vấn lãi suất 6.9%", "Trường stage trong bảng customers được cập nhật, bảng customer_notes thêm mới 1 dòng ghi chú.", "Đạt (Pass)"],
    ], [1800, 2200, 5000, 1000])

    # TC 10: Chat trực tuyến
    add_heading(doc, "3.3.2.10. Kịch bản kiểm thử chức năng Chat trực tuyến", 3)
    add_table(doc, ["Bước thực hiện", "Dữ liệu nhập vào", "Kết quả mong đợi", "Trạng thái thực tế"], [
        ["1. Khách hàng mở khung chat, gửi tin nhắn: 'Xe Camry có sẵn màu đen không?'", "Người gửi: Khách hàng\nNội dung: 'Xe Camry có sẵn màu đen không?'", "Tin nhắn lưu vào chat_messages với is_read = 0. Dashboard nhân viên lập tức nhấp nháy thông báo chat mới.", "Đạt (Pass)"],
        ["2. Nhân viên mở khung chat của khách, gửi tin nhắn trả lời.", "Người gửi: Nhân viên\nNội dung: 'Dạ xe Camry đen đang có sẵn tại showroom ạ'", "Tin nhắn hiển thị tức thì trên khung chat của khách hàng dạng thời gian thực.", "Đạt (Pass)"],
    ], [1800, 2200, 5000, 1000])

    add_heading(doc, "3.4. Đánh giá kết quả", 2)
    add_para(doc, "Trải qua các kịch bản kiểm thử thực tế, toàn bộ hệ thống website bán ô tô trực tuyến showroom TQ Auto đạt được tỷ lệ lỗi phát sinh ở mức dưới 1%. Các luồng dữ liệu nghiệp vụ quan trọng từ đăng ký tư vấn -> ghi nhận lead CRM -> gán nhân viên điều phối lịch hẹn hoạt động trơn tru và có tính đồng bộ thời gian thực.", True)
    add_para(doc, "Sự tách biệt cấu trúc dữ liệu MySQL giữa các bảng liên hệ giúp hạn chế xung đột truy cập ghi dữ liệu cùng lúc. Giao diện responsive hiển thị ổn định trên hầu hết các kích thước màn hình phổ biến. Tuy nhiên, hệ thống vẫn cần cải tiến cơ chế tự động hóa phân chia Lead cho nhân viên theo thuật toán xoay vòng ở phiên bản tiếp theo.", True)
    doc.add_page_break()

    # KẾT LUẬN
    add_heading(doc, "KẾT LUẬN", 1)
    add_para(doc, "Đề tài đã hoàn thành việc xây dựng website bán ô tô trực tuyến dành cho showroom TQ Auto đáp ứng đầy đủ các mục tiêu đề ra ban đầu. Hệ thống cung cấp giải pháp chuyển đổi số thiết thực cho showroom ô tô, giúp loại bỏ các quy trình ghi chép sổ sách thủ công trước đây và thay thế bằng hệ quản trị cơ sở dữ liệu quan hệ đồng bộ.", True)
    add_para(doc, "Thông qua quá trình nghiên cứu và thực thi đồ án tốt nghiệp này, bản thân sinh viên đã nắm vững kiến thức chuyên sâu về lập trình ứng dụng web hiện đại Next.js, cách phân tích thiết kế hệ thống chuẩn hóa bằng ngôn ngữ mô hình hóa UML, thiết kế cơ sở dữ liệu đạt chuẩn tối ưu và thực hiện quy trình kiểm thử phần mềm khoa học.", True)

    # HƯỚNG PHÁT TRIỂN
    add_heading(doc, "HƯỚNG PHÁT TRIỂN", 1)
    add_bullets(doc, [
        "Tích hợp cổng thanh toán trực tuyến (VNPAY/Momo) để khách hàng có thể đặt cọc giữ xe trực tiếp trên website.",
        "Xây dựng module AI thông minh hỗ trợ tư vấn lựa chọn xe tự động dựa trên ngân sách và nhu cầu di chuyển của khách.",
        "Ứng dụng mô hình xem xe 3D tương tác đa chiều (3D Model / 360 Exterior Tour) để nâng cao trải nghiệm xem xe tại nhà.",
    ])

    # TÀI LIỆU THAM KHẢO
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_numbered(doc, [
        "Tài liệu chính thức Next.js Developer Documentation: https://nextjs.org/docs",
        "Tài liệu lập trình React Components: https://react.dev/",
        "Tài liệu hệ quản trị cơ sở dữ liệu MySQL Reference Manual: https://dev.mysql.com/doc/",
        "Giáo trình Phân tích và Thiết kế hệ thống thông tin hướng đối tượng UML - Đại học Công nghiệp Hà Nội.",
        "Tài liệu hướng dẫn triển khai kiểm thử phần mềm (Software Testing Guide).",
        "Mã nguồn dự án TQ Auto trong workspace D:\\web oto - Copy.",
    ])

    doc.core_properties.title = "Báo cáo đồ án tốt nghiệp - Website bán ô tô trực tuyến TQ Auto"
    doc.core_properties.author = "Lê Đình Quốc"
    try:
        doc.save(OUT)
        print(f"Generated complete report at: {OUT}")
    except PermissionError:
        alternative_out = ROOT / "Bao_cao_DATN_TQAuto_LeDinhQuoc_2022603416_v2.docx"
        doc.save(alternative_out)
        print(f"Warning: Primary file was locked (open in MS Word). Saved instead to: {alternative_out}")


if __name__ == "__main__":
    main()
